"""
Word 文档导出器
使用 python-docx 生成 Word 文档
支持 Jinja2 模板和简单占位符替换
"""
from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Pt
from jinja2 import Environment
from src.processors.table_processor import TableProcessor
from src.processors.chart_processor import ChartProcessor
from src.processors.image_processor import ImageProcessor
from src.models.data_models import DataStructure
from src.exporters.base_exporter import BaseExporter


class WordExporter(BaseExporter):
    """
    Word 文档导出器
    支持 Jinja2 模板和简单占位符替换
    """
    
    def __init__(self):
        """初始化 Word 导出器"""
        super().__init__()
        self.table_processor = TableProcessor()
        self.chart_processor = ChartProcessor()
        self.image_processor = ImageProcessor()
        # 创建 jinja2 环境
        self.jinja2_env = Environment()
        # 注册自定义函数和过滤器
        from src.utils.jinja2_filters import JINJA2_GLOBALS, JINJA2_FILTERS
        self.jinja2_env.globals.update(JINJA2_GLOBALS)
        self.jinja2_env.filters.update(JINJA2_FILTERS)
    
    def export(
        self,
        template_path: Optional[Path],
        data: DataStructure,
        output_path: Path,
        watermark: bool = False,
        watermark_text: str = "内部使用，禁止外传",
        watermark_image_path: Optional[str] = None,
        restrict_edit: bool = False,
        restrict_edit_password: Optional[str] = None
    ) -> Path:
        """
        导出 Word 文档
        
        Args:
            template_path: Word 模板文件路径，如果为 None 则使用默认模板
            data: 标准化的数据结构
            output_path: 输出文件路径
        
        Returns:
            生成的文档路径
        """
        # 如果没有模板，使用默认模板生成器
        if template_path is None or not template_path.exists():
            from src.core.default_template_generator import DefaultTemplateGenerator
            doc = DefaultTemplateGenerator.generate_word_template(data)
        else:
            # 填充模板
            doc = self.fill_template(template_path, data)
        
        # 添加水印（如果启用）
        if watermark:
            from src.utils.word_protection import WordProtection
            try:
                if watermark_image_path and Path(watermark_image_path).exists():
                    # 使用图片水印
                    doc = WordProtection.add_image_watermark(doc, watermark_image_path)
                else:
                    # 使用文本水印
                    doc = WordProtection.add_watermark(doc, watermark_text)
            except Exception as e:
                print(f"警告：添加水印失败: {e}")
                import traceback
                traceback.print_exc()
        
        # 限制编辑（如果启用）
        if restrict_edit:
            from src.utils.word_protection import WordProtection
            try:
                doc = WordProtection.restrict_editing(doc, restrict_edit_password)
            except Exception as e:
                print(f"警告：限制编辑设置失败: {e}")
        
        # 保存文件
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        return output_path
    
    def fill_template(
        self,
        template_path: Path,
        data: DataStructure
    ) -> Document:
        """
        填充 Word 模板数据
        
        Args:
            template_path: 模板文件路径
            data: 数据结构
        
        Returns:
            填充后的 Document 对象
        """
        # 加载模板
        doc = Document(str(template_path))
        
        # 检查模板是否包含 Jinja2 语法
        template_text = '\n'.join([para.text for para in doc.paragraphs])
        has_jinja2 = '{%' in template_text or ('{{' in template_text and ('.' in template_text or '|' in template_text))
        
        # 打印模板内容（前500字符）用于调试
        print(f"[DEBUG WordExporter] 模板内容预览（前500字符）: {template_text[:500]}")
        print(f"[DEBUG WordExporter] 模板是否包含Jinja2语法: {has_jinja2}")
        print(f"[DEBUG WordExporter] 数据tables键: {list(data.tables.keys()) if data.tables else '无'}")
        if data.tables:
            for table_name, table_data in data.tables.items():
                print(f"[DEBUG WordExporter] 表格 '{table_name}': {len(table_data) if isinstance(table_data, list) else 'N/A'} 行")
        
        if has_jinja2:
            # 尝试使用 Jinja2 渲染
            print(f"[DEBUG WordExporter] 尝试使用Jinja2渲染模板")
            try:
                result = self._fill_template_with_jinja2(template_path, data)
                # 检查是否成功（如果返回的文档仍然包含占位符，说明失败了）
                result_text = '\n'.join([para.text for para in result.paragraphs])
                # 检查是否还有未处理的占位符（排除已处理的表格和图表占位符）
                has_unprocessed = '{{' in result_text or '{%' in result_text
                # 但表格和图表占位符会在后续处理，所以这里只检查其他占位符
                if not has_unprocessed or ('{{table:' not in result_text and '{{chart:' not in result_text):
                    print(f"[DEBUG WordExporter] Jinja2渲染成功")
                    return result
                else:
                    print(f"[DEBUG WordExporter] Jinja2渲染后仍有占位符，但会继续处理表格和图表")
                    return result  # 即使有占位符，也返回结果，因为后续会处理
            except Exception as e:
                print(f"[DEBUG WordExporter] Jinja2渲染异常: {e}，fallback到简单替换")
                import traceback
                traceback.print_exc()
        
        # 使用原有的简单占位符替换
        # 处理所有段落
        print(f"[DEBUG WordExporter] 使用简单占位符替换，开始填充模板")
        
        # 先处理表格占位符（需要在替换文本之前，因为表格会替换整个段落）
        paragraphs_to_process = list(doc.paragraphs)  # 创建副本，避免迭代时修改
        for paragraph in paragraphs_to_process:
            original_text = paragraph.text
            
            # 处理表格占位符（优先处理，因为表格会替换段落）
            for table_name, table_data in data.tables.items():
                placeholder = f"{{{{table:{table_name}}}}}"
                if placeholder in original_text:
                    print(f"[DEBUG WordExporter] 找到表格占位符: {placeholder}, 表格数据行数: {len(table_data) if isinstance(table_data, list) else 'N/A'}")
                    
                    # 获取表格合并配置（从 data.data 中获取）
                    merge_cells = None
                    if hasattr(data, 'data') and isinstance(data.data, dict):
                        table_merge = data.data.get('table_merge', {})
                        if 'merge_rows' in table_merge and isinstance(table_merge['merge_rows'], list):
                            # 转换合并配置：考虑表头行（+1）
                            merge_rows = []
                            for merge_item in table_merge['merge_rows']:
                                if isinstance(merge_item, dict):
                                    # 表格有表头，所以数据行号需要+1
                                    adjusted_merge = {
                                        'start_row': merge_item.get('start_row', 0) + 1,
                                        'end_row': merge_item.get('end_row', 0) + 1,
                                        'start_col': merge_item.get('start_col', 0),
                                        'end_col': merge_item.get('end_col', 0)
                                    }
                                    merge_rows.append(adjusted_merge)
                            merge_cells = {'merge_rows': merge_rows}
                    
                    # 使用 process_for_word 处理表格（传递合并配置）
                    success = self.table_processor.process_for_word(
                        doc, table_name, table_data, merge_cells=merge_cells
                    )
                    
                    if success:
                        print(f"[DEBUG WordExporter] 表格 '{table_name}' 已成功插入")
                        # 表格已插入，跳过后续处理
                        continue
        
        # 再次遍历段落，处理文本占位符和其他占位符
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            
            # 跳过已处理的表格占位符段落（已被清空）
            if not original_text.strip() and len(doc.tables) > 0:
                # 可能是表格占位符已被处理，跳过
                continue
            
            # 替换文本占位符（包括从tables中提取的数据）
            paragraph.text = self._replace_placeholders_in_text(paragraph.text, data)
            
            # 处理图表占位符
            for chart_name, chart_info in data.charts.items():
                if f"{{{{chart:{chart_name}}}}}" in paragraph.text:
                    chart_type = chart_info.get('type', 'line')
                    chart_data = chart_info.get('data', {})
                    self.chart_processor.process_for_word(
                        doc, chart_name, chart_data, chart_type
                    )
            
            # 处理图片占位符
            # 方式1：从 data.images 字典中获取（已处理的格式）
            for image_name, image_source in data.images.items():
                if f"{{{{image:{image_name}}}}}" in paragraph.text:
                    # 获取图片尺寸（如果原始数据中有）
                    img_width = None
                    img_height = None
                    if hasattr(data, 'data') and isinstance(data.data, dict) and 'images' in data.data:
                        images_array = data.data['images']
                        for img_info in images_array:
                            if isinstance(img_info, dict) and img_info.get('alt') == image_name:
                                img_width = img_info.get('width')
                                img_height = img_info.get('height')
                                break
                    
                    # 转换像素为英寸（假设 96 DPI）
                    width_inches = img_width / 96.0 if img_width else None
                    height_inches = img_height / 96.0 if img_height else None
                    
                    self.image_processor.process_for_word(
                        doc, image_name, image_source,
                        width=width_inches, height=height_inches
                    )
            
            # 方式2：从原始数据中的 images 数组获取（支持 {{images[0].src}} 格式）
            if hasattr(data, 'data') and isinstance(data.data, dict) and 'images' in data.data:
                images_array = data.data['images']
                import re
                # 查找 {{images[数字].src}} 格式的占位符
                image_pattern = r'\{\{images\[(\d+)\]\.src\}\}'
                matches = re.findall(image_pattern, paragraph.text)
                for idx_str in matches:
                    idx = int(idx_str)
                    if idx < len(images_array):
                        img_info = images_array[idx]
                        # 优先使用id字段，如果没有id则使用src字段
                        img_id = img_info.get('id')
                        img_src = img_info.get('src', '')
                        img_alt = img_info.get('alt', f'image_{idx}')
                        # 确定图片源：优先使用ID，否则使用src
                        if img_id is not None:
                            image_source = f'image_id:{img_id}'
                        elif img_src:
                            image_source = img_src
                        else:
                            continue  # 跳过没有图片源的项
                        
                        placeholder = f"{{{{images[{idx}].src}}}}"
                        paragraph.text = paragraph.text.replace(placeholder, '')
                        # 获取图片尺寸
                        img_width = img_info.get('width')
                        img_height = img_info.get('height')
                        width_inches = img_width / 96.0 if img_width else None
                        height_inches = img_height / 96.0 if img_height else None
                        try:
                            self.image_processor.process_for_word(
                                doc, img_alt, image_source,
                                width=width_inches, height=height_inches
                            )
                            print(f"[DEBUG] 插入图片（数组格式）: {img_alt}, 源: {image_source}")
                        except Exception as e:
                            print(f"[WARNING] 插入图片失败: {e}")
        
        # 处理表格中的占位符（如果模板中有表格）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # 替换文本占位符
                    cell.text = self._replace_placeholders_in_text(cell.text, data)
        
        return doc
    
    def _replace_placeholders_in_text(self, text: str, data: DataStructure) -> str:
        """
        替换文本中的占位符，包括从CSV数据中提取的变量
        
        Args:
            text: 包含占位符的文本
            data: 数据结构
        
        Returns:
            替换后的文本
        """
        result = text
        
        # 使用基类的replace_text_placeholder方法
        result = self.replace_text_placeholder(result, data)
        
        # 处理从tables['data']中提取的变量（用于CSV数据）
        if hasattr(data, 'tables') and data.tables:
            tasks_data = data.tables.get('data', [])
            if tasks_data:
                from datetime import datetime
                # 替换常用变量
                result = result.replace('{{project_name}}', str(data.title or '项目名称'))
                result = result.replace('{{report_date}}', datetime.now().strftime('%Y-%m-%d'))
                result = result.replace('{{project_manager}}', '项目经理')
                
                # 计算任务统计
                completed_count = sum(1 for t in tasks_data if 'completed' in str(t.get('Status', '')).lower() or '完成' in str(t.get('Status', '')))
                in_progress_count = sum(1 for t in tasks_data if 'in progress' in str(t.get('Status', '')).lower() or '进行' in str(t.get('Status', '')))
                not_started_count = sum(1 for t in tasks_data if 'not started' in str(t.get('Status', '')).lower() or '未开始' in str(t.get('Status', '')))
                total_count = len(tasks_data)
                
                result = result.replace('{{tasks_list|length}}', str(total_count))
                result = result.replace('{{completed_tasks|length}}', str(completed_count))
                result = result.replace('{{in_progress_tasks|length}}', str(in_progress_count))
                result = result.replace('{{not_started_tasks|length}}', str(not_started_count))
        
        return result
    
    def _fill_template_with_jinja2(
        self,
        template_path: Path,
        data: DataStructure
    ) -> Document:
        """
        使用 Jinja2 渲染 Word 模板
        
        Args:
            template_path: 模板文件路径
            data: 数据结构
        
        Returns:
            填充后的 Document 对象
        """
        # 加载模板文档
        doc = Document(str(template_path))
        
        # 提取所有段落文本
        template_lines = [para.text for para in doc.paragraphs]
        template_text = '\n'.join(template_lines)
        
        # 准备 Jinja2 模板变量（使用原始数据）
        template_vars = {}
        
        # 添加常用Jinja2全局变量（先添加，确保不会被覆盖）
        from datetime import datetime as dt_class, timedelta as td_class
        # 提供now作为可调用函数（用于模板中使用now()）
        template_vars['now'] = lambda: dt_class.now()
        # 提供datetime类（用于模板中使用datetime.now()）
        template_vars['datetime'] = dt_class
        # 提供timedelta类（用于时间计算）
        template_vars['timedelta'] = td_class
        # 提供一个now_func函数（作为now的别名）
        template_vars['now_func'] = lambda: dt_class.now()
        
        # 方案1：不再在代码中硬编码数据转换逻辑
        # 模板应该像HTML模板一样，直接使用 tables['data'] 在Jinja2中处理数据
        # 这样可以实现"换模板不需要改代码"
        # 
        # 注意：为了向后兼容，如果模板仍使用旧变量（tasks_list等），
        # 可以通过检测模板内容来决定是否提供这些变量
        # 但推荐做法是修改模板，使其直接使用 tables['data']
        
        # 如果数据包含原始 JSON 数据（data.data），先展开它
        # 这是关键：将原始JSON数据的所有字段展开到模板变量中，实现"换模板不需要改代码"
        if hasattr(data, 'data') and isinstance(data.data, dict):
            # 直接展开 data.data 中的所有键到模板变量
            # 这样可以访问所有模板需要的变量，如 document、chart_data、images、table_data 等
            for key, value in data.data.items():
                # 优先级：保留已处理的数据
                if key not in template_vars:
                    template_vars[key] = value
        
        # 兼容性处理：如果模板中使用 table_data 变量（如 test3.json）
        if 'table_data' in template_vars and isinstance(template_vars['table_data'], list):
            # 确保 tables['data'] 也指向相同的数据，以便兼容不同模板写法
            if 'tables' not in template_vars or not template_vars['tables']:
                template_vars['tables'] = {'data': template_vars['table_data']}
            elif isinstance(template_vars['tables'], dict) and 'data' not in template_vars['tables']:
                template_vars['tables']['data'] = template_vars['table_data']
        
        # 添加选项变量（用于模板中的条件判断）
        # enable_chart 和 enable_table 从 data.data 中获取，如果没有则默认为 True
        if hasattr(data, 'data') and isinstance(data.data, dict):
            template_vars['enable_chart'] = data.data.get('enable_chart', True)
            template_vars['enable_table'] = data.data.get('enable_table', True)
        else:
            template_vars['enable_chart'] = True
            template_vars['enable_table'] = True
        
        # 添加基本变量
        template_vars['project_name'] = getattr(data, 'title', '项目名称')
        template_vars['report_date'] = dt_class.now().strftime('%Y-%m-%d')
        template_vars['project_manager'] = '项目经理'
        
        # 添加标准字段（最后添加，确保优先级）
        template_vars['title'] = getattr(data, 'title', '')
        template_vars['content'] = getattr(data, 'content', '')
        template_vars['tables'] = getattr(data, 'tables', {})
        template_vars['charts'] = getattr(data, 'charts', {})
        template_vars['images'] = getattr(data, 'images', {})
        
        # 重要：如果原始数据中有 images 数组，也要添加到模板变量中
        # 这样模板可以使用 images[0].alt 这种语法
        if hasattr(data, 'data') and isinstance(data.data, dict) and 'images' in data.data:
            template_vars['images'] = data.data['images']  # 使用原始数组格式
        
        # 为了支持模板中直接使用 data 变量（从 tables['data'] 获取）
        if hasattr(data, 'tables') and data.tables and 'data' in data.tables:
            template_vars['data'] = data.tables['data']
        
        # 提供 today 变量（当前日期）
        template_vars['today'] = dt_class.now().date()
        
        # 两阶段处理：先Jinja2渲染，再处理占位符
        # 阶段1：保护简单的占位符（不包含嵌套Jinja2语法的）
        placeholder_map = {}  # 统一管理所有占位符
        import re
        
        # 处理表格占位符 {{table:name}}（简单格式，不包含嵌套）
        table_pattern = r'\{\{table:(\w+)\}\}(?!\s*\{\{)'
        table_matches = re.findall(table_pattern, template_text)
        for table_name in table_matches:
            placeholder = f"{{{{table:{table_name}}}}}"
            temp_placeholder = f"__TABLE_PLACEHOLDER_{table_name}__"
            placeholder_map[temp_placeholder] = placeholder
            template_text = template_text.replace(placeholder, temp_placeholder)
        
        # 处理简单的图表占位符 {{chart:name}}（不包含嵌套Jinja2的）
        # 注意：不保护 {{chart:{{ chart_data.title }}}} 这种嵌套格式，让Jinja2先渲染
        simple_chart_pattern = r'\{\{chart:([^}]+)\}\}(?!\s*\{\{)'
        chart_matches = re.findall(simple_chart_pattern, template_text)
        for chart_name in chart_matches:
            # 检查是否包含Jinja2变量（如 {{ chart_data.title }}）
            if '{{' in chart_name or '}}' in chart_name:
                # 跳过嵌套的，让Jinja2先渲染
                continue
            placeholder = f"{{{{chart:{chart_name}}}}}"
            temp_placeholder = f"__CHART_PLACEHOLDER_{chart_name.replace(' ', '_')}__"
            placeholder_map[temp_placeholder] = placeholder
            template_text = template_text.replace(placeholder, temp_placeholder)
        
        # 处理简单的图片占位符 {{image:name}}（不包含嵌套的）
        simple_image_pattern = r'\{\{image:([^}]+)\}\}(?!\s*\{\{)'
        image_matches = re.findall(simple_image_pattern, template_text)
        for image_name in image_matches:
            if '{{' in image_name or '}}' in image_name:
                continue
            placeholder = f"{{{{image:{image_name}}}}}"
            temp_placeholder = f"__IMAGE_PLACEHOLDER_{image_name.replace(' ', '_')}__"
            placeholder_map[temp_placeholder] = placeholder
            template_text = template_text.replace(placeholder, temp_placeholder)
        
        # 阶段2：使用 Jinja2 渲染（包括嵌套的占位符）
        try:
            jinja2_template = self.jinja2_env.from_string(template_text)
            rendered_text = jinja2_template.render(**template_vars)
            
            # 恢复所有被保护的占位符
            for temp_placeholder, original_placeholder in placeholder_map.items():
                rendered_text = rendered_text.replace(temp_placeholder, original_placeholder)
            
            # 阶段3：处理Jinja2渲染后新出现的占位符（如 {{chart:{{ chart_data.title }}}} 渲染后变成 {{chart:本月销售趋势}}）
            # 查找所有表格占位符
            table_pattern_final = r'\{\{table:(\w+)\}\}'
            table_matches_final = re.findall(table_pattern_final, rendered_text)
            for table_name in table_matches_final:
                placeholder = f"{{{{table:{table_name}}}}}"
                temp_placeholder = f"__TABLE_PLACEHOLDER_{table_name}__"
                if temp_placeholder not in placeholder_map:
                    placeholder_map[temp_placeholder] = placeholder
                    rendered_text = rendered_text.replace(placeholder, temp_placeholder)
            
            # 查找所有图表占位符（包括动态生成的）
            chart_pattern_final = r'\{\{chart:([^}]+)\}\}'
            chart_matches_final = re.findall(chart_pattern_final, rendered_text)
            for chart_name in chart_matches_final:
                placeholder = f"{{{{chart:{chart_name}}}}}"
                temp_placeholder = f"__CHART_PLACEHOLDER_{chart_name.replace(' ', '_').replace(':', '_')}__"
                if temp_placeholder not in placeholder_map:
                    placeholder_map[temp_placeholder] = placeholder
                    rendered_text = rendered_text.replace(placeholder, temp_placeholder)
            
            # 查找所有图片占位符（包括动态生成的）
            image_pattern_final = r'\{\{image:([^}]+)\}\}'
            image_matches_final = re.findall(image_pattern_final, rendered_text)
            for image_name in image_matches_final:
                placeholder = f"{{{{image:{image_name}}}}}"
                temp_placeholder = f"__IMAGE_PLACEHOLDER_{image_name.replace(' ', '_').replace(':', '_')}__"
                if temp_placeholder not in placeholder_map:
                    placeholder_map[temp_placeholder] = placeholder
                    rendered_text = rendered_text.replace(placeholder, temp_placeholder)
            
        except Exception as e:
            print(f"Jinja2 渲染错误: {e}")
            import traceback
            traceback.print_exc()
            # 渲染失败，抛出异常让调用者处理fallback
            raise e
        
        # 从渲染后的文本创建新文档
        new_doc = self._create_doc_from_rendered_text(rendered_text, doc)
        
        # 处理图片占位符（在渲染后的文档中插入图片）
        # 支持两种格式：
        # 1. {{images[0].alt}} 或 {{images[0].src}} - 数组索引格式
        # 2. {{image:签名}} - 名称格式（使用 alt 作为名称）
        if hasattr(data, 'data') and isinstance(data.data, dict) and 'images' in data.data:
            images_array = data.data['images']
            import re
            
            # 方式1：处理 {{images[0].alt}} 或 {{images[0].src}} 格式
            for paragraph in new_doc.paragraphs:
                text = paragraph.text
                # 查找图片占位符模式：{{images[数字].alt}} 或 {images[数字].alt}（Jinja2渲染后可能只剩一个花括号）
                image_pattern = r'\{*\{*images\[(\d+)\]\.(alt|src)\}*\}*'
                matches = re.findall(image_pattern, text)
                if matches:
                    for idx_str, field in matches:
                        idx = int(idx_str)
                        if idx < len(images_array):
                            img_info = images_array[idx]
                            if field == 'alt':
                                # 替换所有可能的占位符格式（包括单花括号和双花括号）
                                alt_value = img_info.get('alt', '')
                                patterns = [
                                    f"{{{{images[{idx}].alt}}}}",  # {{images[0].alt}}
                                    f"{{images[{idx}].alt}}",      # {images[0].alt}
                                    f"images[{idx}].alt",          # images[0].alt
                                ]
                                for pattern in patterns:
                                    if pattern in paragraph.text:
                                        paragraph.text = paragraph.text.replace(pattern, alt_value)
                                        print(f"[DEBUG] 替换图片alt占位符: {pattern} -> {alt_value}")
                                        break
                            elif field == 'src':
                                # 如果是src，尝试插入图片
                                # 优先使用id字段，如果没有id则使用src字段
                                img_id = img_info.get('id')
                                img_src = img_info.get('src', '')
                                # 确定图片源：优先使用ID，否则使用src
                                if img_id is not None:
                                    image_source = f'image_id:{img_id}'
                                elif img_src:
                                    image_source = img_src
                                else:
                                    continue  # 跳过没有图片源的项
                                
                                # 移除所有可能的占位符格式
                                patterns = [
                                    f"{{{{images[{idx}].src}}}}",
                                    f"{{images[{idx}].src}}",
                                    f"images[{idx}].src",
                                ]
                                for pattern in patterns:
                                    if pattern in paragraph.text:
                                        paragraph.text = paragraph.text.replace(pattern, '')
                                        break
                                # 插入图片
                                try:
                                    # 使用image_processor插入图片
                                    img_alt = img_info.get('alt', f'image_{idx}')
                                    self.image_processor.process_for_word(new_doc, img_alt, image_source)
                                    print(f"[DEBUG] 插入图片: {img_alt}, 源: {image_source}")
                                except Exception as e:
                                    print(f"[WARNING] 插入图片失败: {e}")
                                    import traceback
                                    traceback.print_exc()
            
            # 方式2：处理 {{image:名称}} 格式（使用 alt 作为名称）
            # 构建图片映射：alt -> img_info
            image_map = {}
            for img_info in images_array:
                if isinstance(img_info, dict):
                    img_alt = img_info.get('alt', '')
                    if img_alt:
                        image_map[img_alt] = img_info
            
            # 处理 {{image:名称}} 格式的占位符
            for paragraph in new_doc.paragraphs:
                text = paragraph.text
                # 查找 {{image:名称}} 格式的占位符
                image_name_pattern = r'\{\{image:([^}]+)\}\}'
                matches = re.findall(image_name_pattern, text)
                if matches:
                    for img_name in matches:
                        if img_name in image_map:
                            img_info = image_map[img_name]
                            # 优先使用id字段，如果没有id则使用src字段
                            img_id = img_info.get('id')
                            img_src = img_info.get('src', '')
                            # 确定图片源：优先使用ID，否则使用src
                            if img_id is not None:
                                image_source = f'image_id:{img_id}'
                            elif img_src:
                                image_source = img_src
                            else:
                                continue  # 跳过没有图片源的项
                            
                            # 移除占位符
                            placeholder = f"{{{{image:{img_name}}}}}"
                            paragraph.text = paragraph.text.replace(placeholder, '')
                            # 插入图片
                            try:
                                # 获取图片尺寸（如果指定）
                                img_width = img_info.get('width')
                                img_height = img_info.get('height')
                                # 注意：process_for_word 的 width/height 参数是英寸，需要转换
                                # 如果指定了像素尺寸，转换为英寸（假设 96 DPI）
                                width_inches = None
                                height_inches = None
                                if img_width:
                                    width_inches = img_width / 96.0
                                if img_height:
                                    height_inches = img_height / 96.0
                                
                                self.image_processor.process_for_word(
                                    new_doc, img_name, image_source,
                                    width=width_inches, height=height_inches
                                )
                                print(f"[DEBUG] 插入图片（按名称）: {img_name}, 源: {image_source}")
                                except Exception as e:
                                    print(f"[WARNING] 插入图片失败（按名称）: {e}")
                                    import traceback
                                    traceback.print_exc()
                        else:
                            print(f"[WARNING] 未找到图片: {img_name}，可用图片: {list(image_map.keys())}")
        
        # 处理表格占位符（在渲染后的文档中替换表格）
        # 首先从 data.tables 获取，如果没有，从 data.data 获取
        tables_to_process = {}
        if data.tables:
            tables_to_process.update(data.tables)
        # 如果 data.data 中有 table_data，也添加进去
        if hasattr(data, 'data') and isinstance(data.data, dict):
            if 'table_data' in data.data and isinstance(data.data['table_data'], list):
                tables_to_process['table_data'] = data.data['table_data']
        
        for table_name, table_data in tables_to_process.items():
            if not table_data:
                print(f"[WARNING] 表格 '{table_name}' 数据为空")
                continue
            placeholder = f"{{{{table:{table_name}}}}}"
            # 查找并替换表格占位符
            found = False
            for paragraph in new_doc.paragraphs:
                if placeholder in paragraph.text:
                    # 找到占位符段落，获取其位置
                    para_idx = new_doc.paragraphs.index(paragraph)
                    # 清空段落
                    paragraph.clear()
                    # 使用table_processor插入表格
                    # 注意：process_for_word需要查找占位符，但我们已经清空了段落
                    # 所以需要直接创建表格
                    if isinstance(table_data, list) and len(table_data) > 0:
                        # 获取列名
                        columns = list(table_data[0].keys())
                        num_cols = len(columns)
                        num_rows = len(table_data) + 1  # +1 为表头行
                        
                        # 在段落后插入表格
                        table = new_doc.add_table(rows=num_rows, cols=num_cols)
                        table.style = 'Light Grid Accent 1'
                        
                        # 填充表头
                        header_cells = table.rows[0].cells
                        for i, col_name in enumerate(columns):
                            header_cells[i].text = str(col_name)
                            for para in header_cells[i].paragraphs:
                                for run in para.runs:
                                    run.bold = True
                        
                        # 填充数据
                        for row_idx, row_data in enumerate(table_data, start=1):
                            row_cells = table.rows[row_idx].cells
                            for col_idx, col_name in enumerate(columns):
                                value = row_data.get(col_name, '')
                                row_cells[col_idx].text = str(value) if value is not None else ''
                        
                        # 处理表格合并（从 data.data 中获取 table_merge 配置）
                        if hasattr(data, 'data') and isinstance(data.data, dict):
                            table_merge = data.data.get('table_merge', {})
                            if 'merge_rows' in table_merge and isinstance(table_merge['merge_rows'], list):
                                # 转换合并配置：考虑表头行（+1）
                                merge_rows = []
                                for merge_item in table_merge['merge_rows']:
                                    if isinstance(merge_item, dict):
                                        # 表格有表头，所以数据行号需要+1
                                        adjusted_merge = {
                                            'start_row': merge_item.get('start_row', 0) + 1,
                                            'end_row': merge_item.get('end_row', 0) + 1,
                                            'start_col': merge_item.get('start_col', 0),
                                            'end_col': merge_item.get('end_col', 0)
                                        }
                                        merge_rows.append(adjusted_merge)
                                
                                # 使用 table_processor 的合并方法
                                merge_cells = {'merge_rows': merge_rows}
                                self.table_processor._merge_cells_word(table, merge_cells)
                        
                        print(f"[DEBUG] 表格 '{table_name}' 已插入，数据行数: {len(table_data)}")
                        found = True
                        break
            if not found:
                print(f"[WARNING] 未找到表格占位符: {placeholder}，可用表格: {list(tables_to_process.keys())}")
        
        # 处理图表占位符（在渲染后的文档中插入图表）
        # 首先从 data.charts 获取，如果没有，从 data.data 获取
        charts_to_process = {}
        if data.charts:
            charts_to_process.update(data.charts)
        # 如果 data.data 中有 chart_data，也添加进去
        if hasattr(data, 'data') and isinstance(data.data, dict):
            if 'chart_data' in data.data and isinstance(data.data['chart_data'], dict):
                chart_data = data.data['chart_data']
                chart_name = chart_data.get('title', 'chart_data')
                charts_to_process[chart_name] = chart_data
        
        for chart_name, chart_info in charts_to_process.items():
            placeholder = f"{{{{chart:{chart_name}}}}}"
            # 查找并替换图表占位符
            found = False
            for paragraph in new_doc.paragraphs:
                if placeholder in paragraph.text:
                    # 清空段落
                    paragraph.clear()
                    # 添加图表
                    chart_type = chart_info.get('type', 'line')
                    # chart_info本身就是图表数据，不需要再取data字段
                    chart_data = chart_info
                    try:
                        self.chart_processor.process_for_word(new_doc, chart_name, chart_data, chart_type)
                        print(f"[DEBUG] 图表 '{chart_name}' 已插入，类型: {chart_type}")
                        found = True
                    except Exception as e:
                        print(f"[WARNING] 插入图表失败: {e}")
                        import traceback
                        traceback.print_exc()
                    break
            if not found:
                print(f"[WARNING] 未找到图表占位符: {placeholder}，可用图表: {list(charts_to_process.keys())}")
        
        return new_doc
    
    def _create_doc_from_rendered_text(self, text: str, reference_doc: Document = None) -> Document:
        """
        从渲染后的文本创建 Word 文档
        
        Args:
            text: 渲染后的文本内容
            reference_doc: 参考文档（用于复制样式）
        
        Returns:
            Word 文档对象
        """
        doc = Document()
        
        # 设置默认样式
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        
        lines = text.split('\n')
        in_table = False
        table_rows = []
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            
            # 保留表格和图表占位符，不处理它们（后续会处理）
            if '{{table:' in line_stripped or '{{chart:' in line_stripped:
                # 直接添加占位符段落，不进行任何处理
                doc.add_paragraph(line_stripped)
                continue
            
            if not line_stripped:
                if in_table and table_rows:
                    # 结束表格
                    table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                    for row_idx, row_data in enumerate(table_rows):
                        for col_idx, cell_data in enumerate(row_data):
                            table.rows[row_idx].cells[col_idx].text = str(cell_data)
                    table_rows = []
                    in_table = False
                doc.add_paragraph()
                continue
            
            # 检查是否是表格行（以 | 分隔）
            if '|' in line_stripped and not line_stripped.startswith('#'):
                cells = [cell.strip() for cell in line_stripped.split('|')]
                # 过滤掉空的首尾元素（因为 |a|b| 分割会得到 ['', 'a', 'b', '']）
                cells = [c for c in cells if c]
                # 跳过表格分隔符行（全是 - 的行）
                if cells and all(c.replace('-', '').replace(' ', '') == '' for c in cells):
                    continue
                if cells:
                    table_rows.append(cells)
                    in_table = True
                continue
            
            # 检查是否是标题
            if line_stripped.startswith('#'):
                level = len(line_stripped) - len(line_stripped.lstrip('#'))
                heading_text = line_stripped.lstrip('#').strip()
                doc.add_heading(heading_text, level=min(level, 9))
            elif line_stripped == '---':
                # 忽略分隔线
                continue
            else:
                # 处理段落，将Markdown格式转换为Word格式
                para = doc.add_paragraph()
                # 处理粗体 **文本** 格式
                import re
                text = line_stripped
                # 找到所有 **文本** 模式（使用原始字符串避免转义问题）
                parts = re.split(r'(\*\*[^*]+\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        # 这是粗体文本
                        bold_text = part[2:-2]  # 移除 **
                        run = para.add_run(bold_text)
                        run.bold = True
                    else:
                        # 普通文本
                        para.add_run(part)
        
        # 处理最后的表格
        if in_table and table_rows:
            # 确保所有行的列数一致
            if table_rows:
                max_cols = max(len(row) for row in table_rows)
                # 统一所有行的列数
                for row in table_rows:
                    while len(row) < max_cols:
                        row.append('')
                
                table = doc.add_table(rows=len(table_rows), cols=max_cols)
                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < max_cols:
                            # 处理单元格中的粗体格式
                            cell = table.rows[row_idx].cells[col_idx]
                            cell_text = str(cell_data)
                            # 移除Markdown格式标记
                            cell_text = cell_text.replace('**', '')
                            cell.text = cell_text
        
        return doc
