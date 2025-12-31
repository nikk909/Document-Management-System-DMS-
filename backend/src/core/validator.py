"""
格式校验器
检查文档完整性、链接有效性、样式一致性等
"""
from pathlib import Path
from typing import List, Dict, Any
from src.models.data_models import DataStructure


class Validator:
    """
    格式校验器
    负责检查文档的完整性和正确性
    """
    
    def __init__(self, check_links: bool = False, strict_mode: bool = False):
        """
        初始化校验器
        
        Args:
            check_links: 是否检查链接有效性
            strict_mode: 严格模式（任何错误都导致失败）
        """
        self.check_links = check_links
        self.strict_mode = strict_mode
    
    def validate(
        self,
        document_path: Path,
        data: DataStructure,
        file_format: str
    ) -> List[Dict[str, Any]]:
        """
        执行完整的格式校验
        
        Args:
            document_path: 生成的文档路径
            data: 原始数据结构
            file_format: 文件格式（'word'/'pdf'/'html'）
        
        Returns:
            问题列表，每个问题是一个字典：
            {
                'type': 'error'/'warning',
                'field': '字段名',
                'message': '错误信息'
            }
        """
        problems = []
        
        # 检查文件是否存在
        if not document_path.exists():
            problems.append({
                'type': 'error',
                'field': 'document',
                'message': f'生成的文档不存在: {document_path}'
            })
            return problems
        
        # 检查文件大小
        file_size = document_path.stat().st_size
        if file_size == 0:
            problems.append({
                'type': 'error',
                'field': 'document',
                'message': '生成的文档为空'
            })
        
        # 检查数据填充
        problems.extend(self.validate_data_filling(data, document_path, file_format))
        
        # 检查链接有效性（如果启用）
        if self.check_links:
            problems.extend(self.validate_links(document_path, file_format))
        
        # 检查样式一致性
        problems.extend(self.validate_style_consistency(document_path, file_format))
        
        return problems
    
    def validate_data_filling(
        self,
        data: DataStructure,
        document_path: Path,
        file_format: str
    ) -> List[Dict[str, Any]]:
        """
        检查数据填充是否完整
        检查表格行数是否匹配、必需字段是否填充等
        
        Args:
            data: 原始数据结构
            document_path: 文档路径
            file_format: 文件格式
        
        Returns:
            问题列表
        """
        problems = []
        
        # 检查表格数据（确保 tables 是字典类型）
        if not isinstance(data.tables, dict):
            problems.append({
                'type': 'error',
                'field': 'tables',
                'message': f'表格数据格式错误：期望字典类型，实际为 {type(data.tables).__name__}'
            })
            return problems  # 如果格式错误，直接返回，避免后续错误
        
        for table_name, table_data in data.tables.items():
            if not table_data:
                problems.append({
                    'type': 'warning',
                    'field': f'tables.{table_name}',
                    'message': f'表格数据为空: {table_name}'
                })
            else:
                # 检查表格行数（粗略检查）
                expected_rows = len(table_data)
                # 这里只能做粗略检查，因为无法直接读取文档中的表格行数
                # 实际应用中可能需要解析文档内容
        
        # 检查必需字段
        if not data.title:
            problems.append({
                'type': 'warning',
                'field': 'title',
                'message': '标题为空'
            })
        
        # 检查图表数据（确保 charts 是字典类型）
        if isinstance(data.charts, dict):
            for chart_name, chart_info in data.charts.items():
                # chart_info 就是 chart_data 本身，不是包含 {'data': {...}} 的结构
                if not chart_info or (isinstance(chart_info, dict) and not chart_info):
                    problems.append({
                        'type': 'warning',
                        'field': f'charts.{chart_name}',
                        'message': f'图表数据为空: {chart_name}'
                    })
                elif not isinstance(chart_info, dict):
                    problems.append({
                        'type': 'warning',
                        'field': f'charts.{chart_name}',
                        'message': f'图表数据格式错误: {chart_name}，期望字典类型，实际为 {type(chart_info).__name__}'
                    })
        elif isinstance(data.charts, list):
            problems.append({
                'type': 'warning',
                'field': 'charts',
                'message': f'图表数据格式错误：期望字典类型，实际为列表'
            })
        
        # 检查图片数据（确保 images 是字典类型）
        if isinstance(data.images, dict):
            for image_name, image_source in data.images.items():
                if not image_source:
                    problems.append({
                        'type': 'warning',
                        'field': f'images.{image_name}',
                        'message': f'图片源为空: {image_name}'
                    })
                elif isinstance(image_source, (str, Path)):
                    # 检查是否是URL或Base64格式（这些不需要验证本地文件存在）
                    # 注意：image_source 可能是字符串或 Path 对象，先转换为字符串
                    image_source_str = str(image_source)
                    is_url = (
                        image_source_str.startswith('http://') or 
                        image_source_str.startswith('https://') or
                        '/api/images/' in image_source_str
                    )
                    is_base64 = (
                        image_source_str.startswith('data:image') or
                        image_source_str.startswith('base64:') or
                        image_source_str.startswith('base64,')
                    )
                    
                    if not is_url and not is_base64:
                        # 只有本地文件路径才需要检查文件是否存在
                        image_path = Path(image_source)
                        if not image_path.exists():
                            problems.append({
                                'type': 'error',
                                'field': f'images.{image_name}',
                                'message': f'图片文件不存在: {image_source}'
                            })
        elif isinstance(data.images, list):
            problems.append({
                'type': 'warning',
                'field': 'images',
                'message': f'图片数据格式错误：期望字典类型，实际为列表'
            })
        
        return problems
    
    def validate_links(
        self,
        document_path: Path,
        file_format: str
    ) -> List[Dict[str, Any]]:
        """
        检查文档内的链接有效性
        支持 Word、PDF、HTML 三种格式
        
        Args:
            document_path: 文档路径
            file_format: 文件格式
        
        Returns:
            问题列表
        """
        problems = []
        
        if file_format == 'html':
            # HTML 链接检查
            try:
                content = document_path.read_text(encoding='utf-8')
                import re
                
                # 查找所有链接
                link_pattern = r'href=["\']([^"\']+)["\']'
                links = re.findall(link_pattern, content)
                
                # 检查每个链接
                for link in links:
                    if link.startswith('http://') or link.startswith('https://'):
                        # 外部链接，只做格式检查
                        pass
                    elif link.startswith('#'):
                        # 锚点链接，检查目标是否存在
                        anchor = link[1:]
                        if anchor and f'id="{anchor}"' not in content and f'name="{anchor}"' not in content:
                            problems.append({
                                'type': 'warning',
                                'field': 'links',
                                'message': f'锚点链接目标不存在: {link}'
                            })
                    elif link.startswith('mailto:'):
                        # 邮箱链接，只做格式检查
                        pass
                    else:
                        # 相对路径链接，检查文件是否存在
                        link_path = (document_path.parent / link).resolve()
                        if not link_path.exists():
                            problems.append({
                                'type': 'warning',
                                'field': 'links',
                                'message': f'链接文件不存在: {link}'
                            })
            except Exception as e:
                problems.append({
                    'type': 'warning',
                    'field': 'links',
                    'message': f'检查链接时出错: {e}'
                })
        
        elif file_format == 'word':
            # Word 链接检查
            try:
                from docx import Document
                import re
                from docx.opc.constants import RELATIONSHIP_TYPE as RT
                
                doc = Document(str(document_path))
                
                # 提取所有超链接
                hyperlinks = []
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        if run._element.xml:
                            # 查找超链接元素
                            link_matches = re.findall(
                                r'<w:hyperlink[^>]*r:id="([^"]+)"',
                                run._element.xml
                            )
                            hyperlinks.extend(link_matches)
                
                # 检查链接关系
                if hasattr(doc.part, 'rels'):
                    for rel in doc.part.rels.values():
                        if rel.reltype == RT.HYPERLINK:
                            target = rel.target_ref
                            if target:
                                # 外部链接检查
                                if target.startswith('http://') or target.startswith('https://'):
                                    # 只做格式检查，不实际请求
                                    pass
                                elif target.startswith('#'):
                                    # 内部书签链接，检查书签是否存在
                                    bookmark_name = target[1:]
                                    found_bookmark = False
                                    for paragraph in doc.paragraphs:
                                        if paragraph._element.xml and f'w:name="{bookmark_name}"' in paragraph._element.xml:
                                            found_bookmark = True
                                            break
                                    if not found_bookmark:
                                        problems.append({
                                            'type': 'warning',
                                            'field': 'links',
                                            'message': f'书签链接目标不存在: {target}'
                                        })
                                elif target.startswith('file://'):
                                    # 本地文件链接
                                    file_path = Path(target.replace('file:///', ''))
                                    if not file_path.exists():
                                        problems.append({
                                            'type': 'warning',
                                            'field': 'links',
                                            'message': f'链接文件不存在: {target}'
                                        })
            except ImportError:
                problems.append({
                    'type': 'warning',
                    'field': 'links',
                    'message': 'python-docx 库未安装，无法检查 Word 链接'
                })
            except Exception as e:
                problems.append({
                    'type': 'warning',
                    'field': 'links',
                    'message': f'检查 Word 链接时出错: {e}'
                })
        
        elif file_format == 'pdf':
            # PDF 链接检查
            try:
                import PyPDF2
                
                with open(document_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    
                    # 遍历所有页面
                    for page_num, page in enumerate(pdf_reader.pages):
                        if '/Annots' in page:
                            annotations = page['/Annots']
                            for annotation in annotations:
                                annotation_obj = annotation.get_object()
                                if annotation_obj.get('/Subtype') == '/Link':
                                    # 获取链接目标
                                    action = annotation_obj.get('/A')
                                    if action:
                                        uri = action.get('/URI')
                                        if uri:
                                            link = uri
                                            # 检查外部链接
                                            if link.startswith('http://') or link.startswith('https://'):
                                                # 只做格式检查
                                                pass
                                            elif link.startswith('#'):
                                                # 内部锚点链接
                                                # PDF 内部链接检查较复杂，这里只做简单检查
                                                pass
            except ImportError:
                problems.append({
                    'type': 'warning',
                    'field': 'links',
                    'message': 'PyPDF2 库未安装，无法检查 PDF 链接'
                })
            except Exception as e:
                problems.append({
                    'type': 'warning',
                    'field': 'links',
                    'message': f'检查 PDF 链接时出错: {e}'
                })
        
        return problems
    
    def validate_style_consistency(
        self,
        document_path: Path,
        file_format: str
    ) -> List[Dict[str, Any]]:
        """
        检查样式一致性
        检查字体、页眉页脚等是否统一
        支持 Word、HTML、PDF 格式
        
        Args:
            document_path: 文档路径
            file_format: 文件格式
        
        Returns:
            问题列表
        """
        problems = []
        
        if file_format == 'word':
            # Word 样式检查
            try:
                from docx import Document
                from docx.shared import Pt
                
                doc = Document(str(document_path))
                
                # 收集所有字体信息
                fonts = {}
                font_sizes = {}
                font_colors = {}
                
                # 检查正文段落样式
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        # 字体名称
                        font_name = run.font.name
                        if font_name:
                            fonts[font_name] = fonts.get(font_name, 0) + 1
                        
                        # 字体大小
                        font_size = run.font.size
                        if font_size:
                            if isinstance(font_size, Pt):
                                size_val = font_size.pt
                            else:
                                size_val = str(font_size)
                            font_sizes[size_val] = font_sizes.get(size_val, 0) + 1
                        
                        # 字体颜色
                        font_color = run.font.color
                        if font_color and font_color.rgb:
                            color_val = str(font_color.rgb)
                            font_colors[color_val] = font_colors.get(color_val, 0) + 1
                
                # 检查字体一致性
                if len(fonts) > 3:
                    problems.append({
                        'type': 'warning',
                        'field': 'style',
                        'message': f'检测到 {len(fonts)} 种不同字体，可能存在字体不统一问题。字体列表: {", ".join(fonts.keys()[:5])}'
                    })
                
                if len(font_sizes) > 3:
                    problems.append({
                        'type': 'warning',
                        'field': 'style',
                        'message': f'检测到 {len(font_sizes)} 种不同字体大小，可能存在大小不统一问题'
                    })
                
                # 检查页眉页脚样式
                header_fonts = set()
                footer_fonts = set()
                
                if doc.sections:
                    for section in doc.sections:
                        # 检查页眉
                        header = section.header
                        for paragraph in header.paragraphs:
                            for run in paragraph.runs:
                                if run.font.name:
                                    header_fonts.add(run.font.name)
                        
                        # 检查页脚
                        footer = section.footer
                        for paragraph in footer.paragraphs:
                            for run in paragraph.runs:
                                if run.font.name:
                                    footer_fonts.add(run.font.name)
                
                # 检查页眉页脚字体是否一致
                if header_fonts and footer_fonts:
                    if header_fonts != footer_fonts:
                        problems.append({
                            'type': 'warning',
                            'field': 'style',
                            'message': f'页眉页脚字体不一致。页眉字体: {", ".join(header_fonts)}, 页脚字体: {", ".join(footer_fonts)}'
                        })
                
                # 检查页眉页脚与正文字体是否一致
                main_fonts = set(fonts.keys())
                if header_fonts and main_fonts:
                    if not header_fonts.issubset(main_fonts) and not main_fonts.issubset(header_fonts):
                        problems.append({
                            'type': 'info',
                            'field': 'style',
                            'message': '页眉字体与正文字体不完全一致'
                        })
            
            except ImportError:
                problems.append({
                    'type': 'warning',
                    'field': 'style',
                    'message': 'python-docx 库未安装，无法检查 Word 样式'
                })
            except Exception as e:
                problems.append({
                    'type': 'warning',
                    'field': 'style',
                    'message': f'检查 Word 样式时出错: {e}'
                })
        
        elif file_format == 'html':
            try:
                content = document_path.read_text(encoding='utf-8')
                import re
                
                # 检查是否有内联样式（可能表示样式不统一）
                inline_styles = re.findall(r'style=["\'][^"\']+["\']', content)
                if len(inline_styles) > 10:  # 如果内联样式过多，可能表示样式不统一
                    problems.append({
                        'type': 'warning',
                        'field': 'style',
                        'message': f'检测到 {len(inline_styles)} 个内联样式，可能存在样式不统一问题。建议使用 CSS 类统一管理样式'
                    })
                
                # 检查是否有多个不同的字体定义
                font_families = re.findall(r'font-family:\s*([^;]+)', content)
                unique_fonts = set([f.strip().strip('"\'') for f in font_families])
                if len(unique_fonts) > 5:
                    problems.append({
                        'type': 'warning',
                        'field': 'style',
                        'message': f'检测到 {len(unique_fonts)} 种不同字体，可能存在字体不统一问题'
                    })
            
            except Exception as e:
                problems.append({
                    'type': 'warning',
                    'field': 'style',
                    'message': f'检查 HTML 样式时出错: {e}'
                })
        
        elif file_format == 'pdf':
            # PDF 样式检查较复杂，这里只做基本检查
            # PDF 是通过 HTML 生成的，所以样式检查应该在 HTML 阶段完成
                pass
        
        return problems

