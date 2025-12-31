"""
默认模板生成器
当没有指定模板时，根据数据结构自动生成格式规范的文档
支持智能识别 JSON/CSV 数据，生成简洁易懂的表格格式
"""
from pathlib import Path
from typing import Dict, Any, List, Union
from src.models.data_models import DataStructure
import json


class DefaultTemplateGenerator:
    """
    默认模板生成器
    根据数据结构自动生成格式规范的文档模板
    智能处理 JSON/CSV 数据，生成美观简洁的表格格式
    """
    
    @staticmethod
    def generate_word_template(data: DataStructure) -> Any:
        """
        生成 Word 文档（无模板模式）
        根据数据结构自动创建格式规范的 Word 文档
        
        Args:
            data: 数据结构对象
        
        Returns:
            Document 对象（python-docx）
        """
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # 设置默认样式
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        
        # 1. 添加标题
        if data.title:
            title = doc.add_heading(data.title, level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 2. 添加内容
        if data.content:
            doc.add_paragraph(data.content)
        
        # 3. 处理表格数据（智能识别并创建表格，支持合并单元格）
        has_tables = DefaultTemplateGenerator._add_tables_to_word(doc, data)
        
        # 4. 处理图表数据（支持动态图表生成）
        DefaultTemplateGenerator._add_charts_to_word(doc, data)
        
        # 5. 处理图片数据（支持Base64和本地路径）
        DefaultTemplateGenerator._add_images_to_word(doc, data)
        
        # 6. 处理嵌套的 JSON 数据（如果存在，且表格已处理则跳过重复数据）
        # 避免重复：如果已经有表格数据，不再重复处理 data.data 中的相同数据
        if hasattr(data, 'data') and isinstance(data.data, dict):
            # 如果表格数据为空，才处理 data.data（避免重复生成表格）
            if not has_tables:
                DefaultTemplateGenerator._add_json_data_to_word(doc, data.data, "数据详情")
        
        return doc
    
    @staticmethod
    def _add_tables_to_word(doc: Any, data: DataStructure) -> bool:
        """
        将表格数据添加到 Word 文档（创建真正的表格）
        
        Args:
            doc: Document 对象
            data: 数据结构对象
        
        Returns:
            bool: 是否添加了表格（True 表示已添加表格）
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        has_tables = False
        
        # 处理已提取的表格数据
        for table_name, table_data in data.tables.items():
            if table_data and isinstance(table_data, list) and len(table_data) > 0:
                has_tables = True
                # 添加表格标题
                doc.add_heading(table_name if table_name != 'data' else "数据表格", level=2)
                
                # 获取列名
                if isinstance(table_data[0], dict):
                    columns = list(table_data[0].keys())
                else:
                    # 如果第一行不是字典，使用索引作为列名
                    columns = [f"列{i+1}" for i in range(len(table_data[0]))] if isinstance(table_data[0], (list, tuple)) else ["值"]
                
                if columns:
                    # 创建表格
                    num_rows = len(table_data) + 1  # +1 为表头
                    num_cols = len(columns)
                    
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Light Grid Accent 1'
                    
                    # 填充表头
                    header_cells = table.rows[0].cells
                    for col_idx, col_name in enumerate(columns):
                        header_cells[col_idx].text = str(col_name)
                        # 表头加粗
                        for paragraph in header_cells[col_idx].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                    
                    # 填充数据行
                    for row_idx, row_data in enumerate(table_data, start=1):
                        row_cells = table.rows[row_idx].cells
                        
                        if isinstance(row_data, dict):
                            # 字典类型，按列名提取值
                            for col_idx, col_name in enumerate(columns):
                                value = row_data.get(col_name, '')
                                row_cells[col_idx].text = DefaultTemplateGenerator._format_value(value)
                        elif isinstance(row_data, (list, tuple)):
                            # 列表类型，直接按索引填充
                            for col_idx in range(min(num_cols, len(row_data))):
                                row_cells[col_idx].text = DefaultTemplateGenerator._format_value(row_data[col_idx])
                        else:
                            # 简单值
                            row_cells[0].text = DefaultTemplateGenerator._format_value(row_data)
                    
                    # 处理表格合并单元格（如果数据中有table_merge配置）
                    if hasattr(data, 'data') and isinstance(data.data, dict) and 'table_merge' in data.data:
                        merge_config = data.data['table_merge']
                        if 'merge_rows' in merge_config:
                            DefaultTemplateGenerator._merge_table_cells_word(table, merge_config['merge_rows'])
                    
                    doc.add_paragraph()  # 添加空行
        
        return has_tables
    
    @staticmethod
    def _merge_table_cells_word(table: Any, merge_rows: List[Dict[str, int]]):
        """
        合并Word表格单元格
        
        Args:
            table: Word表格对象
            merge_rows: 合并配置列表，格式：[{"start_row": 0, "end_row": 1, "start_col": 0, "end_col": 0}]
        """
        for merge_config in merge_rows:
            start_row = merge_config.get('start_row', 0) + 1  # +1因为表头行
            end_row = merge_config.get('end_row', 0) + 1
            start_col = merge_config.get('start_col', 0)
            end_col = merge_config.get('end_col', 0)
            
            if start_row < len(table.rows) and end_row < len(table.rows):
                if start_col == end_col and start_row <= end_row:
                    # 合并同一列的多个行
                    for row_idx in range(start_row, end_row + 1):
                        if row_idx < len(table.rows):
                            cell1 = table.rows[start_row].cells[start_col]
                            cell2 = table.rows[row_idx].cells[start_col]
                            if cell1 != cell2:
                                cell1.merge(cell2)
                elif start_row == end_row and start_col <= end_col:
                    # 合并同一行的多个列
                    if start_row < len(table.rows):
                        row = table.rows[start_row]
                        for col_idx in range(start_col + 1, end_col + 1):
                            if col_idx < len(row.cells):
                                cell1 = row.cells[start_col]
                                cell2 = row.cells[col_idx]
                                if cell1 != cell2:
                                    cell1.merge(cell2)
    
    @staticmethod
    def _add_charts_to_word(doc: Any, data: DataStructure):
        """
        将图表添加到Word文档
        
        Args:
            doc: Document对象
            data: 数据结构对象
        """
        if not data.charts:
            return
        
        from src.processors.chart_processor import ChartProcessor
        chart_processor = ChartProcessor()
        
        for chart_name, chart_info in data.charts.items():
            try:
                chart_type = chart_info.get('type', 'line')
                # 生成图表并添加到文档
                chart_path = chart_processor.generate_chart(chart_info, chart_type)
                
                # 添加图表标题
                chart_title = chart_info.get('title', chart_name)
                doc.add_heading(chart_title, level=2)
                
                # 添加图表图片
                from docx.shared import Inches
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(str(chart_path), width=Inches(6))
                
                # 删除临时文件
                if chart_path.exists():
                    chart_path.unlink()
                
                doc.add_paragraph()  # 添加空行
            except Exception as e:
                print(f"处理图表 {chart_name} 时出错: {e}")
    
    @staticmethod
    def _add_images_to_word(doc: Any, data: DataStructure):
        """
        将图片添加到Word文档
        
        Args:
            doc: Document对象
            data: 数据结构对象
        """
        if not data.images:
            return
        
        from src.processors.image_processor import ImageProcessor
        from docx.shared import Inches
        image_processor = ImageProcessor()
        
        # 如果images是字典，直接处理
        if isinstance(data.images, dict):
            for image_name, image_source in data.images.items():
                try:
                    # 加载图片数据
                    image_data = image_processor.load_image(image_source)
                    
                    # 添加图片标题
                    doc.add_heading(image_name, level=3)
                    
                    # 添加图片
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()
                    
                    # 写入临时文件
                    temp_path = Path.cwd() / f"temp_image_{image_name}.png"
                    temp_path.write_bytes(image_data)
                    
                    try:
                        # 获取图片尺寸并添加到文档
                        img_size = image_processor.get_image_size(image_data)
                        if img_size[0] > 0:
                            width = Inches(min(img_size[0] / 96, 6))  # 最大6英寸
                            run.add_picture(str(temp_path), width=width)
                        else:
                            run.add_picture(str(temp_path), width=Inches(4))
                    finally:
                        if temp_path.exists():
                            temp_path.unlink()
                    
                    doc.add_paragraph()  # 添加空行
                except Exception as e:
                    print(f"处理图片 {image_name} 时出错: {e}")
        
    @staticmethod
    def _format_value(value: Any) -> str:
        """
        格式化值，使其更易读
        缺失值（None、NaN、空字符串）统一显示为 "null"
        
        Args:
            value: 待格式化的值
        
        Returns:
            格式化后的字符串
        """
        import math
        import pandas as pd
        
        # 检查是否为 None
        if value is None:
            return 'null'
        
        # 检查是否为数组类型（numpy/pandas），如果是数组则转换为列表
        try:
            import numpy as np
            if isinstance(value, (np.ndarray, pd.Series)):
                return json.dumps(value.tolist(), ensure_ascii=False, separators=(',', ':'))
        except (ImportError, AttributeError):
            pass
        
        # 检查是否为 NaN（pandas 或 numpy）
        try:
            if isinstance(value, float) and math.isnan(value):
                return 'null'
        except (ValueError, TypeError):
            pass
        
        try:
            if isinstance(value, (pd._libs.missing.NAType, type(pd.NA))):
                return 'null'
        except (AttributeError, TypeError):
            pass
        
        # 检查是否为 pandas NaN（需要先检查是否为标量，避免数组判断错误）
        try:
            import numpy as np
            if hasattr(pd, 'isna'):
                # 如果是标量值，才使用 pd.isna
                if not isinstance(value, (np.ndarray, pd.Series, list, dict, tuple)):
                    if pd.isna(value):
                        return 'null'
        except (ImportError, AttributeError, ValueError, TypeError):
            pass
        
        # 检查是否为字符串 "nan" 或 "NaN"
        if isinstance(value, str) and value.lower() in ('nan', 'none', ''):
            return 'null'
        elif isinstance(value, bool):
            return '是' if value else '否'
        elif isinstance(value, (dict, list)):
            # 对于复杂类型，使用 JSON 格式但压缩
            return json.dumps(value, ensure_ascii=False, separators=(',', ':'))
        else:
            return str(value)
    
    @staticmethod
    def _add_json_data_to_word(doc: Any, json_data: Dict[str, Any], section_title: str, level: int = 2):
        """
        将 JSON 数据添加到 Word 文档（智能识别列表并创建表格）
        
        Args:
            doc: Document 对象
            json_data: JSON 数据字典
            section_title: 章节标题
            level: 标题级别
        """
        if not json_data:
            return
        
        # 添加章节标题
        doc.add_heading(section_title, level=level)
        
        # 遍历 JSON 数据
        for key, value in json_data.items():
            if isinstance(value, dict):
                # 嵌套字典，递归处理
                DefaultTemplateGenerator._add_json_data_to_word(doc, value, str(key), level + 1)
            elif isinstance(value, list) and value:
                # 列表数据，尝试创建表格
                if isinstance(value[0], dict):
                    # 字典列表，创建表格
                    DefaultTemplateGenerator._create_table_from_list(doc, value, str(key))
                elif isinstance(value[0], (list, tuple)):
                    # 列表的列表，创建表格
                    DefaultTemplateGenerator._create_table_from_list(doc, value, str(key))
                else:
                    # 简单值列表，创建简单的表格
                    DefaultTemplateGenerator._create_simple_list_table(doc, value, str(key))
            else:
                # 简单值或空列表
                if isinstance(value, list) and len(value) == 0:
                    doc.add_paragraph(f"{key}: (空列表)")
                elif isinstance(value, (dict, list)):
                    pretty_value = json.dumps(value, ensure_ascii=False, indent=2)
                    doc.add_paragraph(f"{key}:")
                    for line in pretty_value.splitlines():
                        doc.add_paragraph(line)
                else:
                    doc.add_paragraph(f"{key}: {DefaultTemplateGenerator._format_value(value)}")
        
        doc.add_paragraph()  # 添加空行
    
    @staticmethod
    def _create_table_from_list(doc: Any, data_list: List[Any], table_title: str):
        """
        从列表数据创建表格
        
        Args:
            doc: Document 对象
            data_list: 数据列表
            table_title: 表格标题
        """
        if not data_list:
            return
        
        # 添加表格标题（小标题）
        doc.add_paragraph(table_title, style='Heading 3')
        
        # 确定列数和列名
        first_item = data_list[0]
        if isinstance(first_item, dict):
            columns = list(first_item.keys())
        elif isinstance(first_item, (list, tuple)):
            columns = [f"列{i+1}" for i in range(len(first_item))]
        else:
            # 简单值，创建单列表格
            columns = ["值"]
        
        if not columns:
            return
        
        # 创建表格
        num_rows = len(data_list) + 1  # +1 为表头
        num_cols = len(columns)
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light Grid Accent 1'
        
        # 填充表头
        header_cells = table.rows[0].cells
        for col_idx, col_name in enumerate(columns):
            header_cells[col_idx].text = str(col_name)
            for paragraph in header_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # 填充数据
        for row_idx, item in enumerate(data_list, start=1):
            row_cells = table.rows[row_idx].cells
            
            if isinstance(item, dict):
                for col_idx, col_name in enumerate(columns):
                    value = item.get(col_name, '')
                    row_cells[col_idx].text = DefaultTemplateGenerator._format_value(value)
            elif isinstance(item, (list, tuple)):
                for col_idx in range(min(num_cols, len(item))):
                    row_cells[col_idx].text = DefaultTemplateGenerator._format_value(item[col_idx])
            else:
                row_cells[0].text = DefaultTemplateGenerator._format_value(item)
        
        doc.add_paragraph()  # 添加空行
    
    @staticmethod
    def _create_simple_list_table(doc: Any, data_list: List[Any], table_title: str):
        """
        从简单值列表创建表格
        
        Args:
            doc: Document 对象
            data_list: 简单值列表
            table_title: 表格标题
        """
        if not data_list:
            return
        
        doc.add_paragraph(table_title, style='Heading 3')
        
        # 创建两列表格：序号和值
        num_rows = len(data_list) + 1
        table = doc.add_table(rows=num_rows, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = "序号"
        header_cells[1].text = "值"
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # 数据行
        for row_idx, value in enumerate(data_list, start=1):
            row_cells = table.rows[row_idx].cells
            row_cells[0].text = str(row_idx)
            row_cells[1].text = DefaultTemplateGenerator._format_value(value)
        
        doc.add_paragraph()
    
    @staticmethod
    def generate_html_template(data: DataStructure) -> str:
        """
        生成 HTML 文档（无模板模式）
        根据数据结构自动创建格式规范的 HTML 文档
        
        Args:
            data: 数据结构对象
        
        Returns:
            HTML 内容字符串
        """
        html_parts = []
        
        # HTML 头部（改进的样式）
        html_parts.append("""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>""" + (data.title or "文档") + """</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        body {
            font-family: "Microsoft YaHei", "SimSun", "宋体", Arial, sans-serif;
            line-height: 1.6;
            margin: 0;
            padding: 40px;
            background-color: #f5f5f5;
            color: #333;
        }
        .container {
            max-width: 1200px;
            margin: 0 auto;
            background-color: #fff;
            padding: 40px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }
        h1 {
            text-align: center;
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 15px;
            margin-bottom: 30px;
            font-size: 2em;
        }
        h2 {
            color: #34495e;
            margin-top: 30px;
            margin-bottom: 15px;
            border-left: 4px solid #3498db;
            padding-left: 15px;
            font-size: 1.5em;
        }
        h3 {
            color: #555;
            margin-top: 20px;
            margin-bottom: 10px;
            font-size: 1.2em;
        }
        .content {
            margin: 20px 0;
            padding: 15px;
            background-color: #f9f9f9;
            border-left: 4px solid #3498db;
            border-radius: 4px;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
            background-color: #fff;
            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        }
        th, td {
            border: 1px solid #ddd;
            padding: 12px 15px;
            text-align: left;
            font-size: 14px;
        }
        th {
            background-color: #3498db;
            color: white;
            font-weight: bold;
            text-align: center;
        }
        tr:nth-child(even) {
            background-color: #f9f9f9;
        }
        tr:hover {
            background-color: #e8f4f8;
        }
        td {
            word-wrap: break-word;
            max-width: 300px;
        }
        .json-section {
            margin: 20px 0;
            padding: 15px;
            background-color: #f5f5f5;
            border-radius: 5px;
            border-left: 3px solid #95a5a6;
        }
        .json-key {
            font-weight: bold;
            color: #2980b9;
        }
        .info-item {
            margin: 10px 0;
            padding: 8px 0;
            border-bottom: 1px solid #eee;
        }
        .info-label {
            font-weight: bold;
            color: #555;
            display: inline-block;
            width: 150px;
        }
        .info-value {
            color: #333;
        }
        pre {
            background-color: #f4f4f4;
            padding: 10px;
            border-radius: 4px;
            overflow-x: auto;
            font-size: 12px;
        }
    </style>
</head>
<body>
    <div class="container">""")
        
        # 标题
        if data.title:
            html_parts.append(f"<h1>{data.title}</h1>")
        
        # 内容
        if data.content:
            html_parts.append(f'<div class="content">{data.content}</div>')
        
        # 表格数据（改进：创建真正的 HTML 表格）
        has_tables = DefaultTemplateGenerator._add_tables_to_html(html_parts, data)
                
        # 图表数据（支持动态图表生成）
        DefaultTemplateGenerator._add_charts_to_html(html_parts, data)
        
        # 图片数据（支持Base64和本地路径）
        DefaultTemplateGenerator._add_images_to_html(html_parts, data)
        
        # JSON 数据（智能识别并创建表格）
        # 避免重复：如果已经有表格数据，不再重复处理 data.data 中的相同数据
        if hasattr(data, 'data') and isinstance(data.data, dict):
            # 如果表格数据为空，才处理 data.data（避免重复生成表格）
            if not has_tables:
                html_parts.append(DefaultTemplateGenerator._add_json_data_to_html(data.data, "数据详情"))
        
        # HTML 尾部
        html_parts.append("    </div>\n</body>\n</html>")
        
        return "\n".join(html_parts)
    
    @staticmethod
    def _add_tables_to_html(html_parts: List[str], data: DataStructure) -> bool:
        """
        将表格数据添加到 HTML（创建真正的表格）
        
        Args:
            html_parts: HTML 部分列表
            data: 数据结构对象
        
        Returns:
            bool: 是否添加了表格（True 表示已添加表格）
        """
        has_tables = False
        
        for table_name, table_data in data.tables.items():
            if table_data and isinstance(table_data, list) and len(table_data) > 0:
                has_tables = True
                # 表格标题
                display_name = table_name if table_name != 'data' else "数据表格"
                html_parts.append(f"        <h2>{display_name}</h2>")
                
                # 获取列名
                first_row = table_data[0]
                if isinstance(first_row, dict):
                    columns = list(first_row.keys())
                elif isinstance(first_row, (list, tuple)):
                    columns = [f"列{i+1}" for i in range(len(first_row))]
                else:
                    columns = ["值"]
                
                if columns:
                    # 创建表格
                    html_parts.append("        <table>")
                    # 表头
                    html_parts.append("            <thead>")
                    html_parts.append("                <tr>")
                    for col_name in columns:
                        html_parts.append(f"                    <th>{col_name}</th>")
                    html_parts.append("                </tr>")
                    html_parts.append("            </thead>")
                    # 表体
                    html_parts.append("            <tbody>")
                    for row_data in table_data:
                        html_parts.append("                <tr>")
                        if isinstance(row_data, dict):
                            for col_name in columns:
                                value = row_data.get(col_name, '')
                                html_parts.append(f"                    <td>{DefaultTemplateGenerator._escape_html(DefaultTemplateGenerator._format_value(value))}</td>")
                        elif isinstance(row_data, (list, tuple)):
                            for idx in range(len(columns)):
                                value = row_data[idx] if idx < len(row_data) else ''
                                html_parts.append(f"                    <td>{DefaultTemplateGenerator._escape_html(DefaultTemplateGenerator._format_value(value))}</td>")
                        else:
                            html_parts.append(f"                    <td colspan=\"{len(columns)}\">{DefaultTemplateGenerator._escape_html(DefaultTemplateGenerator._format_value(row_data))}</td>")
                        html_parts.append("                </tr>")
                    html_parts.append("            </tbody>")
                    html_parts.append("        </table>")
        
        return has_tables
    
    @staticmethod
    def _escape_html(text: str) -> str:
        """
        HTML 转义
        
        Args:
            text: 待转义的文本
        
        Returns:
            转义后的文本
        """
        if not isinstance(text, str):
            text = str(text)
        return (text.replace('&', '&amp;')
                   .replace('<', '&lt;')
                   .replace('>', '&gt;')
                   .replace('"', '&quot;')
                   .replace("'", '&#39;'))
    
    @staticmethod
    def _add_json_data_to_html(json_data: Dict[str, Any], section_title: str, level: int = 2) -> str:
        """
        将 JSON 数据添加到 HTML（智能识别列表并创建表格）
        
        Args:
            json_data: JSON 数据字典
            section_title: 章节标题
            level: 标题级别
        
        Returns:
            HTML 字符串
        """
        html_parts = []
        
        if not json_data:
            return ""
        
        html_parts.append(f"        <h{level}>{section_title}</h{level}>")
        html_parts.append('        <div class="json-section">')
        
        for key, value in json_data.items():
            if isinstance(value, dict):
                # 嵌套字典，递归处理
                html_parts.append(DefaultTemplateGenerator._add_json_data_to_html(value, str(key), level + 1))
            elif isinstance(value, list) and value:
                # 列表数据，尝试创建表格
                if isinstance(value[0], dict):
                    # 字典列表，创建表格
                    html_parts.append(DefaultTemplateGenerator._create_html_table_from_list(value, str(key)))
                elif isinstance(value[0], (list, tuple)):
                    # 列表的列表，创建表格
                    html_parts.append(DefaultTemplateGenerator._create_html_table_from_list(value, str(key)))
                else:
                    # 简单值列表，创建简单表格
                    html_parts.append(DefaultTemplateGenerator._create_simple_html_list_table(value, str(key)))
            elif isinstance(value, list) and len(value) == 0:
                # 空列表
                html_parts.append(f'            <div class="info-item"><span class="info-label">{key}:</span><span class="info-value">(空列表)</span></div>')
            else:
                # 简单值
                formatted_value = DefaultTemplateGenerator._format_value(value)
                if isinstance(value, (dict, list)):
                    # 复杂类型，使用代码块显示
                    pretty_value = json.dumps(value, ensure_ascii=False, indent=2)
                    html_parts.append(f'            <p><span class="json-key">{key}:</span></p>')
                    html_parts.append("            <pre>")
                    html_parts.append(DefaultTemplateGenerator._escape_html(pretty_value))
                    html_parts.append("            </pre>")
                else:
                    html_parts.append(f'            <div class="info-item"><span class="info-label">{key}:</span><span class="info-value">{DefaultTemplateGenerator._escape_html(formatted_value)}</span></div>')
        
        html_parts.append("        </div>")
        
        return "\n".join(html_parts)

    @staticmethod
    def _create_html_table_from_list(data_list: List[Any], table_title: str) -> str:
        """
        从列表数据创建 HTML 表格
        
        Args:
            data_list: 数据列表
            table_title: 表格标题
        
        Returns:
            HTML 字符串
        """
        if not data_list:
            return ""
        
        html_parts = []
        html_parts.append(f"            <h3>{table_title}</h3>")
        
        # 确定列数和列名
        first_item = data_list[0]
        if isinstance(first_item, dict):
            columns = list(first_item.keys())
        elif isinstance(first_item, (list, tuple)):
            columns = [f"列{i+1}" for i in range(len(first_item))]
        else:
            columns = ["值"]
        
        if not columns:
            return ""
        
        # 创建表格
        html_parts.append("            <table>")
        # 表头
        html_parts.append("                <thead>")
        html_parts.append("                    <tr>")
        for col_name in columns:
            html_parts.append(f"                        <th>{col_name}</th>")
        html_parts.append("                    </tr>")
        html_parts.append("                </thead>")
        # 表体
        html_parts.append("                <tbody>")
        for item in data_list:
            html_parts.append("                    <tr>")
            if isinstance(item, dict):
                for col_name in columns:
                    value = item.get(col_name, '')
                    html_parts.append(f"                        <td>{DefaultTemplateGenerator._escape_html(DefaultTemplateGenerator._format_value(value))}</td>")
            elif isinstance(item, (list, tuple)):
                for idx in range(len(columns)):
                    value = item[idx] if idx < len(item) else ''
                    html_parts.append(f"                        <td>{DefaultTemplateGenerator._escape_html(DefaultTemplateGenerator._format_value(value))}</td>")
            else:
                html_parts.append(f"                        <td colspan=\"{len(columns)}\">{DefaultTemplateGenerator._escape_html(DefaultTemplateGenerator._format_value(item))}</td>")
            html_parts.append("                    </tr>")
        html_parts.append("                </tbody>")
        html_parts.append("            </table>")
        
        return "\n".join(html_parts)

    @staticmethod
    def _create_simple_html_list_table(data_list: List[Any], table_title: str) -> str:
        """
        从简单值列表创建 HTML 表格
        
        Args:
            data_list: 简单值列表
            table_title: 表格标题
        
        Returns:
            HTML 字符串
        """
        if not data_list:
            return ""
        
        html_parts = []
        html_parts.append(f"            <h3>{table_title}</h3>")
        html_parts.append("            <table>")
        html_parts.append("                <thead>")
        html_parts.append("                    <tr>")
        html_parts.append("                        <th>序号</th>")
        html_parts.append("                        <th>值</th>")
        html_parts.append("                    </tr>")
        html_parts.append("                </thead>")
        html_parts.append("                <tbody>")
        for idx, value in enumerate(data_list, 1):
            html_parts.append("                    <tr>")
            html_parts.append(f"                        <td>{idx}</td>")
            html_parts.append(f"                        <td>{DefaultTemplateGenerator._escape_html(DefaultTemplateGenerator._format_value(value))}</td>")
            html_parts.append("                    </tr>")
        html_parts.append("                </tbody>")
        html_parts.append("            </table>")
        
        return "\n".join(html_parts)

    @staticmethod
    def _add_charts_to_html(html_parts: List[str], data: DataStructure):
        """
        将图表添加到HTML
        
        Args:
            html_parts: HTML部分列表
            data: 数据结构对象
        """
        if not data.charts:
            return
        
        from src.processors.chart_processor import ChartProcessor
        chart_processor = ChartProcessor()
        
        for chart_name, chart_info in data.charts.items():
            try:
                chart_type = chart_info.get('type', 'line')
                # 生成图表并转换为Base64
                base64_str = chart_processor.generate_chart_base64(chart_info, chart_type)
                
                # 添加图表标题
                chart_title = chart_info.get('title', chart_name)
                html_parts.append(f"        <h2>{chart_title}</h2>")
                
                # 添加图表图片
                html_parts.append(f'        <img src="data:image/png;base64,{base64_str}" alt="{chart_title}" style="max-width: 100%; height: auto; margin: 20px 0;" />')
                html_parts.append("")  # 添加空行
            except Exception as e:
                print(f"处理图表 {chart_name} 时出错: {e}")
    
    @staticmethod
    def _add_images_to_html(html_parts: List[str], data: DataStructure):
        """
        将图片添加到HTML
        
        Args:
            html_parts: HTML部分列表
            data: 数据结构对象
        """
        if not data.images:
            return
        
        from src.processors.image_processor import ImageProcessor
        image_processor = ImageProcessor()
        
        # 如果images是字典，直接处理
        if isinstance(data.images, dict):
            for image_name, image_source in data.images.items():
                try:
                    # 加载图片数据
                    image_data = image_processor.load_image(image_source)
                    
                    # 转换为Base64
                    import base64
                    base64_str = base64.b64encode(image_data).decode('utf-8')
                    
                    # 检测图片格式
                    img_format = 'png'
                    if isinstance(image_source, Path):
                        ext = image_source.suffix.lower()
                        if ext in ['.jpg', '.jpeg']:
                            img_format = 'jpeg'
                        elif ext == '.gif':
                            img_format = 'gif'
                    elif isinstance(image_source, str) and image_source.startswith('base64:'):
                        # base64:前缀，假设是PNG
                        img_format = 'png'
                    
                    # 添加图片标题
                    html_parts.append(f"        <h3>{image_name}</h3>")
                    
                    # 添加图片
                    html_parts.append(f'        <img src="data:image/{img_format};base64,{base64_str}" alt="{image_name}" style="max-width: 100%; height: auto; margin: 10px 0;" />')
                    html_parts.append("")  # 添加空行
                except Exception as e:
                    print(f"处理图片 {image_name} 时出错: {e}")


