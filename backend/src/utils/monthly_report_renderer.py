"""
月度报表渲染器
用于将sample-1.json格式的数据渲染为三种格式的报表
支持表格合并、图表生成、图片插入
"""
import base64
from pathlib import Path
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def render_docx_monthly_report(data: Dict[str, Any], output_path: Path) -> Path:
    """
    渲染DOCX格式的月度报表
    
    Args:
        data: 包含document、table_data、table_merge、chart_data、images的数据字典
        output_path: 输出文件路径
    
    Returns:
        生成的文档路径
    
    说明：
    - 使用python-docx的add_table和merge_cells实现表格合并
    - 使用ChartProcessor生成图表并作为图片插入
    - 使用ImageProcessor处理Base64和本地路径图片
    """
    from src.processors.chart_processor import ChartProcessor
    from src.processors.image_processor import ImageProcessor
    
    doc = Document()
    
    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 1. 添加标题
    document_info = data.get('document', {})
    title = document_info.get('title', '月度报表')
    date = document_info.get('date', '')
    
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if date:
        date_para = doc.add_paragraph(date)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para_format = date_para.paragraph_format
        date_para_format.space_after = Pt(12)
    
    doc.add_paragraph()  # 空行
    
    # 2. 添加表格（支持合并单元格）
    table_data = data.get('table_data', [])
    if table_data:
        # 获取列名
        columns = list(table_data[0].keys())
        num_rows = len(table_data)
        num_cols = len(columns)
        
        # 创建表格
        table = doc.add_table(rows=num_rows + 1, cols=num_cols)  # +1 for header
        table.style = 'Light Grid Accent 1'  # Word专业型样式
        
        # 添加表头
        header_cells = table.rows[0].cells
        for col_idx, col_name in enumerate(columns):
            header_cells[col_idx].text = col_name
            header_cells[col_idx].paragraphs[0].runs[0].font.bold = True
        
        # 填充数据
        for row_idx, row_data in enumerate(table_data, start=1):
            row_cells = table.rows[row_idx].cells
            for col_idx, col_name in enumerate(columns):
                value = row_data.get(col_name, '')
                row_cells[col_idx].text = str(value)
        
        # 处理表格合并
        table_merge = data.get('table_merge', {})
        if 'merge_rows' in table_merge:
            for merge in table_merge['merge_rows']:
                start_row = merge.get('start_row', 0) + 1  # +1因为表头
                end_row = merge.get('end_row', 0) + 1
                start_col = merge.get('start_col', 0)
                end_col = merge.get('end_col', 0)
                
                if (start_row < len(table.rows) and end_row < len(table.rows) and
                    start_col < num_cols and end_col < num_cols):
                    # 合并单元格
                    if start_row <= end_row and start_col <= end_col:
                        cell1 = table.rows[start_row].cells[start_col]
                        for r in range(start_row, end_row + 1):
                            for c in range(start_col, end_col + 1):
                                if r != start_row or c != start_col:
                                    cell2 = table.rows[r].cells[c]
                                    cell1.merge(cell2)
        
        doc.add_paragraph()  # 空行
    
    # 3. 添加图表
    chart_data = data.get('chart_data')
    if chart_data:
        chart_processor = ChartProcessor()
        
        try:
            chart_title = chart_data.get('title', '图表')
            doc.add_heading(chart_title, level=2)
            
            # 生成图表
            chart_type = chart_data.get('type', 'line')
            chart_path = chart_processor.generate_chart(chart_data, chart_type)
            
            # 插入图表图片
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(str(chart_path), width=Inches(6))
            
            # 删除临时文件
            if chart_path.exists():
                chart_path.unlink()
            
            doc.add_paragraph()  # 空行
        except Exception as e:
            print(f"处理图表时出错: {e}")
    
    # 4. 添加图片（在文档末尾右侧）
    images = data.get('images', [])
    if images:
        image_processor = ImageProcessor()
        
        # 创建右对齐段落
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        for img_info in images:
            try:
                img_src = img_info.get('src', '')
                img_alt = img_info.get('alt', '')
                img_width = img_info.get('width', 200)
                img_height = img_info.get('height', 50)
                
                # 加载图片
                image_data = image_processor.load_image(img_src)
                
                # 写入临时文件
                temp_path = Path.cwd() / f"temp_image_{img_alt}.png"
                temp_path.write_bytes(image_data)
                
                try:
                    # 添加到段落
                    run = para.add_run()
                    run.add_picture(str(temp_path), width=Inches(img_width / 96))
                    
                    # 添加图片说明
                    if img_alt:
                        run.add_break()
                        run = para.add_run(img_alt)
                        run.font.size = Pt(9)
                finally:
                    if temp_path.exists():
                        temp_path.unlink()
                
                # 添加空格分隔
                para.add_run('  ')
            except Exception as e:
                print(f"处理图片 {img_info.get('alt', '')} 时出错: {e}")
    
    # 保存文档
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    
    return output_path

