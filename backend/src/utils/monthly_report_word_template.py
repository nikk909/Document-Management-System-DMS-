"""
月度报表Word模板渲染器
根据sample-1.json数据生成符合模板结构的Word文档
"""
from pathlib import Path
from typing import Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime


def render_monthly_report_word(data: Dict[str, Any], output_path: Path) -> Path:
    """
    渲染月度报表Word文档（按照用户提供的模板结构）
    
    Args:
        data: 包含document、table_data、table_merge、chart_data、images的数据字典
        output_path: 输出文件路径
    
    Returns:
        生成的文档路径
    """
    from src.processors.chart_processor import ChartProcessor
    from src.processors.image_processor import ImageProcessor
    
    doc = Document()
    
    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)
    
    # 1. 文档标题和日期
    document_info = data.get('document', {})
    title = document_info.get('title', '月度报表')
    date = document_info.get('date', '')
    
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if date:
        date_para = doc.add_paragraph(f'报告日期：{date}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para_format = date_para.paragraph_format
        date_para_format.space_after = Pt(12)
    
    doc.add_paragraph()  # 空行
    
    # 2. 表格部分：部门销售数据，支持合并
    table_data = data.get('table_data', [])
    if table_data:
        doc.add_heading('部门销售汇总', level=2)
        
        # 获取列名
        columns = list(table_data[0].keys())
        num_rows = len(table_data)
        num_cols = len(columns)
        
        # 创建表格
        table = doc.add_table(rows=num_rows + 1, cols=num_cols)
        table.style = 'Light Grid Accent 1'  # Word专业型样式
        
        # 添加表头
        header_cells = table.rows[0].cells
        for col_idx, col_name in enumerate(columns):
            cell = header_cells[col_idx]
            cell.text = col_name if col_name != '销售额' else f'{col_name} (元)'
            # 设置表头加粗
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # 填充数据
        for row_idx, row_data in enumerate(table_data, start=1):
            row_cells = table.rows[row_idx].cells
            for col_idx, col_name in enumerate(columns):
                value = row_data.get(col_name, '')
                if col_name == '增长率':
                    # 格式化增长率为百分比
                    if isinstance(value, (int, float)):
                        row_cells[col_idx].text = f'{value * 100:.1f}%'
                    else:
                        row_cells[col_idx].text = str(value)
                elif col_name == '销售额':
                    # 格式化销售额为千位分隔符
                    if isinstance(value, (int, float)):
                        row_cells[col_idx].text = f'{value:,.0f}'
                    else:
                        row_cells[col_idx].text = str(value)
                else:
                    row_cells[col_idx].text = str(value)
        
        # 处理表格合并
        table_merge = data.get('table_merge', {})
        if 'merge_rows' in table_merge:
            for merge in table_merge['merge_rows']:
                start_row = merge.get('start_row', 0) + 1  # +1因为表头
                end_row = merge.get('end_row', 0) + 1
                start_col = merge.get('start_col', 0)
                end_col = merge.get('end_col', 0)
                
                if (start_row < len(table.rows) and end_row < len(table.rows) and
                    start_col < num_cols and end_col < num_cols):
                    # 合并单元格
                    if start_row <= end_row and start_col <= end_col:
                        cell1 = table.rows[start_row].cells[start_col]
                        for r in range(start_row, end_row + 1):
                            for c in range(start_col, end_col + 1):
                                if r != start_row or c != start_col:
                                    cell2 = table.rows[r].cells[c]
                                    cell1.merge(cell2)
        
        doc.add_paragraph()  # 空行
    
    # 3. 图表部分：销售趋势线图
    chart_data = data.get('chart_data')
    if chart_data:
        chart_title_text = chart_data.get('title', '本月销售趋势')
        doc.add_heading(chart_title_text, level=2)
        
        x_label = chart_data.get('x_label', '周次')
        y_label = chart_data.get('y_label', '销售额')
        chart_desc = doc.add_paragraph(f'X轴：{x_label} | Y轴：{y_label}')
        
        # 创建图表数据表格（用于显示数据点）
        series_list = chart_data.get('series', [])
        labels = chart_data.get('labels', [])
        
        if series_list and labels:
            # 创建表格：1行表头 + 数据行
            chart_table = doc.add_table(rows=len(labels) + 1, cols=1 + len(series_list))
            chart_table.style = 'Light List Accent 1'
            
            # 表头
            chart_header = chart_table.rows[0].cells
            chart_header[0].text = x_label
            for cell in chart_header:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            for idx, serie in enumerate(series_list, start=1):
                chart_header[idx].text = serie.get('name', f'系列{idx}')
                for paragraph in chart_header[idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # 数据行
            for i, label in enumerate(labels):
                row_cells = chart_table.rows[i + 1].cells
                row_cells[0].text = str(label)
                
                for idx, serie in enumerate(series_list, start=1):
                    points = serie.get('points', serie.get('data', []))
                    if i < len(points):
                        value = points[i]
                        if isinstance(value, (int, float)):
                            row_cells[idx].text = f'{value:,.0f}'
                        else:
                            row_cells[idx].text = str(value)
            
            # 图表说明文字
            note_para = doc.add_paragraph('注：以上数据可用于生成线图，显示本月销售趋势上升。')
            note_para_format = note_para.paragraph_format
            note_para_format.space_before = Pt(6)
            
            # 尝试生成实际的图表图片并插入
            try:
                chart_processor = ChartProcessor()
                chart_type = chart_data.get('type', 'line')
                chart_path = chart_processor.generate_chart(chart_data, chart_type)
                
                # 插入图表图片
                chart_img_para = doc.add_paragraph()
                chart_img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = chart_img_para.add_run()
                run.add_picture(str(chart_path), width=Inches(6))
                
                # 删除临时文件
                if chart_path.exists():
                    chart_path.unlink()
            except Exception as e:
                print(f"生成图表图片时出错: {e}")
        
        doc.add_paragraph()  # 空行
    
    # 4. 图像部分
    images = data.get('images', [])
    if images:
        doc.add_heading('附件图像', level=2)
        
        image_processor = ImageProcessor()
        
        for img_info in images:
            img_src = img_info.get('src', '')
            img_alt = img_info.get('alt', '')
            img_width = img_info.get('width', 200)
            img_height = img_info.get('height', 50)
            
            # 处理图片路径（如果是相对路径，转换为绝对路径）
            if isinstance(img_src, str) and not img_src.startswith(('data:', 'base64:')):
                # 处理以 / 开头的路径
                if img_src.startswith('/'):
                    img_src = img_src[1:]  # 去掉开头的 /
                # 尝试多个可能的路径
                possible_paths = [
                    Path(img_src),
                    Path.cwd() / img_src,
                    Path(__file__).parent.parent.parent / img_src,
                ]
                # 找到第一个存在的路径
                found_path = None
                for path in possible_paths:
                    if path.exists():
                        found_path = path
                        break
                if found_path:
                    img_src = str(found_path)
                else:
                    # 如果所有路径都不存在，记录错误
                    img_src = None
            
            try:
                if img_src is None:
                    raise FileNotFoundError(f"图片文件不存在: {img_info.get('src', '')}")
                
                # 加载图片
                image_data = image_processor.load_image(img_src)
                
                # 将图片转换为PNG格式（Word支持PNG/JPG，但不支持WebP）
                from PIL import Image
                import io
                
                try:
                    # 使用PIL打开图片（支持各种格式）
                    img = Image.open(io.BytesIO(image_data))
                    
                    # 如果是RGBA模式的图片，转换为RGB（JPEG不支持透明通道）
                    if img.mode in ('RGBA', 'LA', 'P'):
                        # 创建白色背景
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        if img.mode == 'P':
                            img = img.convert('RGBA')
                        background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
                        img = background
                    elif img.mode != 'RGB':
                        img = img.convert('RGB')
                    
                    # 保存为PNG格式
                    temp_path = Path.cwd() / f"temp_image_{img_alt}_{hash(img_src)}.png"
                    img.save(temp_path, 'PNG')
                except Exception as e:
                    # 如果转换失败，尝试直接写入（可能已经是支持的格式）
                    print(f"图片格式转换失败，尝试直接写入: {e}")
                    temp_path = Path.cwd() / f"temp_image_{img_alt}_{hash(img_src)}.png"
                    temp_path.write_bytes(image_data)
                
                try:
                    # 添加图片（居中）
                    img_para = doc.add_paragraph()
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = img_para.add_run()
                    run.add_picture(str(temp_path), width=Inches(img_width / 96))
                    
                    # 添加图片说明
                    if img_alt:
                        label_para = doc.add_paragraph(img_alt)
                        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        label_para_format = label_para.paragraph_format
                        label_para_format.space_after = Pt(12)
                finally:
                    if temp_path.exists():
                        temp_path.unlink()
            except Exception as e:
                # 当图片加载失败时，在文档中添加错误提示和说明
                error_msg = f"处理图片 {img_alt} 时出错: {str(e)}"
                print(error_msg)
                
                # 在文档中添加图片占位符框（模拟图片区域）
                placeholder_para = doc.add_paragraph()
                placeholder_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 创建一个带边框的占位符
                placeholder_run = placeholder_para.add_run(f"[图片占位符: {img_alt}]")
                placeholder_run.font.color.rgb = RGBColor(200, 200, 200)  # 浅灰色
                placeholder_run.font.size = Pt(10)
                
                # 添加图片说明（必须显示，无论图片是否加载成功）
                if img_alt:
                    label_para = doc.add_paragraph(img_alt)
                    label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    label_run = label_para.runs[0]
                    label_run.font.size = Pt(11)
                    label_para_format = label_para.paragraph_format
                    label_para_format.space_after = Pt(12)
                else:
                    # 即使没有alt，也添加一个空行保持格式
                    doc.add_paragraph()
    
    # 5. 报告生成时间
    time_para = doc.add_paragraph(f'报告生成于 {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    time_para_format = time_para.paragraph_format
    time_para_format.space_before = Pt(12)
    time_run = time_para.runs[0]
    time_run.font.size = Pt(12)
    time_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # 保存文档
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    
    return output_path

