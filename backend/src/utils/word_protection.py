"""
Word 文档保护工具
支持限制编辑和水印功能
"""
from pathlib import Path
from typing import Optional, Union
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class WordProtection:
    """
    Word 文档保护工具
    支持限制编辑和水印功能
    """
    
    @staticmethod
    def restrict_editing(
        doc: Document,
        password: Optional[str] = None,
        allow_only_comments: bool = False,
        allow_only_form_filling: bool = False,
        allow_only_read_only: bool = True
    ) -> Document:
        """
        限制 Word 文档编辑
        
        Args:
            doc: Document 对象
            password: 保护密码（可选）
            allow_only_comments: 仅允许批注
            allow_only_form_filling: 仅允许填写表单
            allow_only_read_only: 仅允许只读（默认）
        
        Returns:
            处理后的 Document 对象
        """
        try:
            # 获取文档的 settings 部分
            settings = doc.settings
            
            # 创建文档保护元素
            # 注意：python-docx 不直接支持限制编辑，需要通过 XML 操作
            # 这里我们通过添加文档属性来实现只读保护
            
            # 方法1：通过修改文档属性设置只读建议
            # 这会在打开文档时提示用户以只读模式打开
            if allow_only_read_only:
                # 设置文档为只读建议
                # 注意：真正的限制编辑需要使用 COM 对象或修改 settings.xml
                # 这里我们通过添加文档属性来实现
                pass  # python-docx 限制，无法直接设置限制编辑
            
            # 方法2：通过添加文档保护标记（在 settings.xml 中）
            # 这需要直接操作 XML
            try:
                # 获取 settings 的 XML 元素
                settings_element = settings._element
                
                # 创建文档保护元素
                # w:documentProtection 元素
                protection = OxmlElement('w:documentProtection')
                protection.set(qn('w:enforcement'), '1')
                
                if allow_only_read_only:
                    protection.set(qn('w:edit'), 'readOnly')
                elif allow_only_comments:
                    protection.set(qn('w:edit'), 'comments')
                elif allow_only_form_filling:
                    protection.set(qn('w:edit'), 'forms')
                
                if password:
                    # 注意：密码需要加密，这里简化处理
                    # 实际应用中需要使用 Word 的加密算法
                    protection.set(qn('w:cryptProviderType'), 'rsaFull')
                    protection.set(qn('w:cryptAlgorithmClass'), 'hash')
                    protection.set(qn('w:cryptAlgorithmType'), 'typeAny')
                    protection.set(qn('w:cryptAlgorithmSid'), '4')
                    protection.set(qn('w:hash'), password)  # 简化处理，实际需要加密
                
                # 添加到 settings
                settings_element.append(protection)
                
            except Exception as e:
                # 如果 XML 操作失败，使用简单方法
                print(f"警告：无法设置文档保护（XML操作失败）: {e}")
                # 可以通过添加水印提示用户文档受保护
                WordProtection.add_watermark(doc, "只读文档，禁止编辑")
        
        except Exception as e:
            print(f"警告：限制编辑设置失败: {e}")
        
        return doc
    
    @staticmethod
    def add_watermark(
        doc: Document,
        text: str = "内部使用，禁止外传",
        font_size: int = 80,
        color: RGBColor = RGBColor(192, 192, 192),
        angle: int = -45
    ) -> Document:
        """
        为 Word 文档添加水印（使用VML在页眉中添加背景水印）
        
        Args:
            doc: Document 对象
            text: 水印文本
            font_size: 字体大小
            color: 颜色（默认浅灰色）
            angle: 旋转角度（默认-45度）
        
        Returns:
            处理后的 Document 对象
        """
        print(f"[DEBUG] WordProtection.add_watermark() 被调用，水印文本: {text}, 文档节数: {len(doc.sections)}")
        
        try:
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls, qn
            
            # 为每个节添加水印
            for section_idx, section in enumerate(doc.sections):
                print(f"[DEBUG] 处理第 {section_idx + 1} 节的水印")
                
                # 获取页眉
                header = section.header
                header.is_linked_to_previous = False
                
                # 清除现有页眉内容
                for para in list(header.paragraphs):
                    para._element.getparent().remove(para._element)
                
                # 创建VML水印XML（使用Word的标准水印格式）
                # 这个XML定义了一个居中的旋转文本形状，z-index为负值，使其在正文背景显示
                watermark_xml = f'''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office">
                    <w:pPr>
                        <w:pStyle w:val="Header"/>
                    </w:pPr>
                    <w:r>
                        <w:rPr>
                            <w:noProof/>
                        </w:rPr>
                        <w:pict>
                            <v:shapetype id="_x0000_t136" coordsize="21600,21600" o:spt="136" path="m@7,l@8,m@5,l@6,m@3,l@4,e">
                                <v:formulas>
                                    <v:f eqn="if lineDrawn pixelLineWidth 0"/>
                                    <v:f eqn="sum @0 1 0"/>
                                    <v:f eqn="sum 0 0 @1"/>
                                    <v:f eqn="prod @2 1 2"/>
                                    <v:f eqn="prod @3 21600 pixelWidth"/>
                                    <v:f eqn="prod @3 21600 pixelHeight"/>
                                    <v:f eqn="sum @0 0 1"/>
                                    <v:f eqn="prod @6 1 2"/>
                                    <v:f eqn="prod @7 21600 pixelWidth"/>
                                    <v:f eqn="prod @7 21600 pixelHeight"/>
                                </v:formulas>
                                <v:path textpathok="t" o:connecttype="rect"/>
                                <o:lock v:ext="edit" aspectratio="t"/>
                                <v:textpath on="t" fitshape="t"/>
                            </v:shapetype>
                            <v:shape id="WordWatermark" type="#_x0000_t136" style="position:absolute;margin-left:0;margin-top:0;width:412.4pt;height:137.45pt;z-index:-251656064;mso-wrap-edited:f;mso-position-horizontal:center;mso-position-horizontal-relative:margin;mso-position-vertical:center;mso-position-vertical-relative:margin;rotation:{angle}" filled="t" stroked="f">
                                <v:fill color="silver" opacity="0.3"/>
                                <v:textpath style="font-family:&quot;宋体&quot;;font-size:{font_size}pt" string="{text}"/>
                            </v:shape>
                        </w:pict>
                    </w:r>
                </w:p>'''
                
                try:
                    # 解析并添加水印XML
                    watermark_element = parse_xml(watermark_xml)
                    header._element.append(watermark_element)
                    print(f"[DEBUG] 第 {section_idx + 1} 节VML水印添加成功")
                except Exception as e:
                    print(f"[WARN] VML水印添加失败: {e}，使用文本水印")
                    # 备用：在页眉中添加文本
                    para = header.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run(text)
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = color
                    run.font.name = '宋体'
            
            print(f"[DEBUG] 所有水印添加完成")
                
        except Exception as e:
            print(f"[WARN] 添加水印失败: {e}")
            import traceback
            traceback.print_exc()
            
            # 最后备用：在文档正文开头添加水印文本
            try:
                if doc.paragraphs:
                    first_para = doc.paragraphs[0]
                    watermark_para = first_para.insert_paragraph_before()
                else:
                    watermark_para = doc.add_paragraph()
                
                watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = watermark_para.add_run(f"【{text}】")
                run.font.size = Pt(font_size)
                run.font.color.rgb = color
                run.font.name = '宋体'
                print(f"[DEBUG] 使用文档正文备用方法添加水印")
            except Exception as e2:
                print(f"[WARN] 备用方法也失败: {e2}")
        
        return doc
    
    @staticmethod
    def add_image_watermark(
        doc: Document,
        image_path: Union[Path, str],
        opacity: float = 0.3,
        angle: int = -45
    ) -> Document:
        """
        为 Word 文档添加图片水印
        
        Args:
            doc: Document 对象
            image_path: 水印图片路径
            opacity: 透明度（0.0-1.0，默认0.3）
            angle: 旋转角度（默认-45度）
        
        Returns:
            处理后的 Document 对象
        """
        try:
            from docx.oxml import parse_xml
            from docx.shared import Inches
            from PIL import Image
            import base64
            import io
            
            image_path = Path(image_path)
            if not image_path.exists():
                raise FileNotFoundError(f"水印图片不存在: {image_path}")
            
            # 读取图片并转换为base64
            with open(image_path, 'rb') as f:
                image_data = f.read()
            
            # 获取图片尺寸
            img = Image.open(io.BytesIO(image_data))
            img_width, img_height = img.size
            
            # 转换图片为base64（Word使用内嵌图片）
            base64_image = base64.b64encode(image_data).decode('utf-8')
            
            # 为每个节添加图片水印
            for section in doc.sections:
                header = section.header
                
                # 创建图片水印段落
                paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                paragraph.clear()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 添加图片（使用run.add_picture）
                run = paragraph.add_run()
                try:
                    # 将图片保存到临时文件以便add_picture使用
                    import tempfile
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
                        tmp_path = tmp_file.name
                        tmp_file.write(image_data)
                    
                    # 添加图片到run
                    run.add_picture(tmp_path, width=Inches(6))  # 默认宽度6英寸
                    
                    # 清理临时文件
                    import os
                    try:
                        os.unlink(tmp_path)
                    except:
                        pass
                except Exception as e:
                    print(f"警告：添加图片水印失败: {e}")
                    import traceback
                    traceback.print_exc()
        except Exception as e:
            print(f"警告：添加图片水印失败: {e}")
            import traceback
            traceback.print_exc()
        
        return doc
    
    @staticmethod
    def add_watermark_to_pdf(
        pdf_path: Union[Path, str],
        text: str = "内部使用，禁止外传",
        output_path: Optional[Union[Path, str]] = None,
        image_path: Optional[Union[Path, str]] = None,
        angle: int = 45
    ) -> Path:
        """
        为 PDF 文档添加水印
        
        Args:
            pdf_path: PDF 文件路径
            text: 水印文本（当image_path为None时使用）
            output_path: 输出文件路径（如果为 None，覆盖原文件）
            image_path: 水印图片路径（如果提供，使用图片水印；否则使用文本水印）
            angle: 旋转角度（默认45度）
        
        Returns:
            处理后的 PDF 文件路径
        """
        pdf_path = Path(pdf_path)
        if output_path is None:
            output_path = pdf_path
        else:
            output_path = Path(output_path)
        
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter, A4
            from PyPDF2 import PdfReader, PdfWriter
            from PIL import Image as PILImage
            import io
            
            # 读取原始 PDF
            reader = PdfReader(str(pdf_path))
            writer = PdfWriter()
            
            # 如果提供了图片路径，使用图片水印，否则使用文本水印
            use_image_watermark = image_path is not None and Path(image_path).exists()
            
            # 为每一页添加水印
            for page_num, page in enumerate(reader.pages):
                # 创建水印 PDF
                packet = io.BytesIO()
                can = canvas.Canvas(packet, pagesize=A4)
                
                if use_image_watermark:
                    # 图片水印
                    try:
                        img = PILImage.open(image_path)
                        # 获取图片尺寸
                        img_width, img_height = img.size
                        # 计算缩放比例（使图片适合页面）
                        page_width, page_height = A4
                        scale = min(page_width * 0.6 / img_width, page_height * 0.6 / img_height)
                        scaled_width = img_width * scale
                        scaled_height = img_height * scale
                        
                        # 居中放置并旋转
                        can.saveState()
                        can.translate(page_width / 2, page_height / 2)
                        can.rotate(angle)
                        # 绘制图片（使用透明度）
                        can.drawImage(str(image_path), -scaled_width/2, -scaled_height/2, 
                                    width=scaled_width, height=scaled_height, 
                                    mask='auto', preserveAspectRatio=True)
                        can.restoreState()
                    except Exception as e:
                        print(f"警告：图片水印失败，使用文本水印: {e}")
                        # fallback到文本水印
                        can.setFont("Helvetica-Bold", 50)
                        can.setFillColorRGB(0.7, 0.7, 0.7, alpha=0.3)
                        can.saveState()
                        can.translate(A4[0] / 2, A4[1] / 2)
                        can.rotate(angle)
                        can.drawString(-200, 0, text)
                        can.restoreState()
                else:
                    # 文本水印
                    can.setFont("Helvetica-Bold", 50)
                    can.setFillColorRGB(0.7, 0.7, 0.7, alpha=0.3)  # 浅灰色，半透明
                    
                    # 旋转并添加水印文本
                    can.saveState()
                    can.translate(A4[0] / 2, A4[1] / 2)
                    can.rotate(angle)
                    can.drawString(-200, 0, text)
                    can.restoreState()
                
                can.save()
                packet.seek(0)
                
                # 读取水印 PDF
                watermark_pdf = PdfReader(packet)
                watermark_page = watermark_pdf.pages[0]
                
                # 合并水印和原页面
                page.merge_page(watermark_page)
                writer.add_page(page)
            
            # 保存带水印的 PDF
            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            return output_path
            
        except ImportError:
            raise ImportError(
                "PDF 水印需要安装 reportlab 和 PyPDF2\n"
                "安装命令: pip install reportlab PyPDF2"
            )
        except Exception as e:
            raise RuntimeError(f"PDF 水印添加失败: {e}")

