"""
表格处理器
处理表格数据的填充，支持自动匹配列数、合并单元格等
"""
import json
import pandas as pd
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Inches
from docx.table import Table


class TableProcessor:
    """
    表格处理器
    负责将数据填充到表格中，支持 Word 和 HTML/PDF 格式
    """
    
    def __init__(self):
        """初始化表格处理器"""
        pass
    
    @staticmethod
    def _format_value(value: Any) -> str:
        """
        格式化值，使其更易读
        将 None, NaN, 空字符串统一显示为 "null"
        
        Args:
            value: 待格式化的值
            
        Returns:
            格式化后的字符串
        """
        if value is None or (isinstance(value, float) and pd.isna(value)) or (isinstance(value, str) and value.strip() == ''):
            return 'null'
        elif isinstance(value, bool):
            return '是' if value else '否'
        elif isinstance(value, (dict, list)):
            return json.dumps(value, ensure_ascii=False, separators=(',', ':'))
        else:
            return str(value)
    
    def process_for_word(
        self,
        doc: Document,
        placeholder: str,
        table_data: List[Dict[str, Any]],
        merge_cells: Optional[Dict[str, List[int]]] = None
    ) -> bool:
        """
        为 Word 文档处理表格
        查找占位符 {{table:name}}，替换为实际表格
        
        Args:
            doc: Word 文档对象（python-docx Document）
            placeholder: 占位符名称（如 'data'）
            table_data: 表格数据，列表格式，每个元素是一个字典
            merge_cells: 合并单元格配置（可选）：
                {
                    'row': [0, 1],  # 要合并的行
                    'col': [0, 1]   # 要合并的列
                }
        
        Returns:
            是否成功处理
        """
        if not table_data:
            return False
        
        # 获取所有列名（从第一条数据中提取）
        columns = list(table_data[0].keys())
        num_cols = len(columns)
        num_rows = len(table_data) + 1  # +1 为表头行
        
        # 遍历所有段落，查找占位符
        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text
            if f"{{{{table:{placeholder}}}}}" in paragraph_text:
                # 找到占位符，创建表格
                # 先删除占位符文本（清空段落）
                paragraph.clear()
                
                # 在段落位置插入表格
                # 注意：python-docx 中，表格是插入在段落之后的
                # 我们需要在段落后添加表格
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Light Grid Accent 1'  # 设置表格样式
                
                # 填充表头
                header_cells = table.rows[0].cells
                for i, col_name in enumerate(columns):
                    header_cells[i].text = str(col_name)
                    # 设置表头加粗
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # 填充数据
                for row_idx, row_data in enumerate(table_data, start=1):
                    if row_idx >= num_rows:
                        break  # 防止索引越界
                    row_cells = table.rows[row_idx].cells
                    for col_idx, col_name in enumerate(columns):
                        if col_idx >= num_cols:
                            break  # 防止索引越界
                        value = row_data.get(col_name, '')
                        row_cells[col_idx].text = TableProcessor._format_value(value)
                
                # 处理合并单元格
                if merge_cells:
                    self._merge_cells_word(table, merge_cells)
                
                return True
        
        return False
    
    def process_for_html(
        self,
        template_content: str,
        placeholder: str,
        table_data: List[Dict[str, Any]],
        merge_cells: Optional[Dict[str, List[int]]] = None
    ) -> str:
        """
        为 HTML/PDF 模板处理表格
        将占位符替换为 HTML 表格
        
        Args:
            template_content: HTML 模板内容
            placeholder: 占位符名称
            table_data: 表格数据
            merge_cells: 合并单元格配置（可选）
        
        Returns:
            处理后的 HTML 内容
        """
        if not table_data:
            return template_content.replace(f"{{{{table:{placeholder}}}}}", "")
        
        # 获取列名
        columns = list(table_data[0].keys())
        
        # 生成 HTML 表格
        html_table = '<table border="1" cellpadding="5" cellspacing="0" style="border-collapse: collapse; width: 100%;">\n'
        
        # 表头
        html_table += '  <thead>\n    <tr>\n'
        for col_name in columns:
            html_table += f'      <th>{col_name}</th>\n'
        html_table += '    </tr>\n  </thead>\n'
        
        # 表体
        html_table += '  <tbody>\n'
        for row_data in table_data:
            html_table += '    <tr>\n'
            for col_name in columns:
                value = row_data.get(col_name, '')
                formatted_value = TableProcessor._format_value(value)
                html_table += f'      <td>{formatted_value}</td>\n'
            html_table += '    </tr>\n'
        html_table += '  </tbody>\n'
        
        html_table += '</table>\n'
        
        # 替换占位符
        return template_content.replace(f"{{{{table:{placeholder}}}}}", html_table)
    
    def _merge_cells_word(self, table: Table, merge_cells: Dict[str, Any]):
        """
        在 Word 表格中合并单元格
        
        Args:
            table: Word 表格对象
            merge_cells: 合并配置，支持两种格式：
                1. 旧格式：{'row': [rows], 'col': [cols]}
                2. 新格式：{'merge_rows': [{'start_row': 0, 'end_row': 1, 'start_col': 0, 'end_col': 0}]}
        """
        # 支持新格式：merge_rows
        if 'merge_rows' in merge_cells and isinstance(merge_cells['merge_rows'], list):
            for merge_item in merge_cells['merge_rows']:
                if isinstance(merge_item, dict):
                    start_row = merge_item.get('start_row', 0)
                    end_row = merge_item.get('end_row', 0)
                    start_col = merge_item.get('start_col', 0)
                    end_col = merge_item.get('end_col', 0)
                    
                    # 确保索引有效
                    if (start_row < len(table.rows) and end_row < len(table.rows) and
                        start_col >= 0 and end_col >= 0):
                        
                        # 合并指定区域
                        start_cell = table.rows[start_row].cells[start_col]
                        end_cell = table.rows[end_row].cells[end_col]
                        
                        # 如果 start 和 end 是同一个单元格，跳过
                        if start_cell == end_cell:
                            continue
                        
                        try:
                            # 合并从 start_cell 到 end_cell 的所有单元格
                            # python-docx 的 merge 方法会合并从起始单元格到目标单元格的矩形区域
                            start_cell.merge(end_cell)
                        except Exception as e:
                            print(f"[WARNING] 合并单元格失败 (行{start_row}-{end_row}, 列{start_col}-{end_col}): {e}")
        
        # 支持旧格式：row 和 col（向后兼容）
        if 'row' in merge_cells and len(merge_cells['row']) >= 2:
            rows = sorted(merge_cells['row'])
            if rows[0] < len(table.rows) and rows[1] < len(table.rows):
                for col_idx in range(len(table.columns)):
                    cell1 = table.rows[rows[0]].cells[col_idx]
                    cell2 = table.rows[rows[1]].cells[col_idx]
                    try:
                        cell1.merge(cell2)
                    except Exception as e:
                        print(f"[WARNING] 合并行失败: {e}")
        
        # 合并列
        if 'col' in merge_cells and len(merge_cells['col']) >= 2:
            cols = sorted(merge_cells['col'])
            if cols[0] >= 0 and cols[1] < len(table.columns):
                for row_idx in range(len(table.rows)):
                    cell1 = table.rows[row_idx].cells[cols[0]]
                    cell2 = table.rows[row_idx].cells[cols[1]]
                    try:
                        cell1.merge(cell2)
                    except Exception as e:
                        print(f"[WARNING] 合并列失败: {e}")
    
    def auto_match_columns(
        self,
        table_data: List[Dict[str, Any]],
        expected_columns: List[str]
    ) -> List[Dict[str, Any]]:
        """
        自动匹配列数
        如果数据列数与期望不一致，自动调整
        
        Args:
            table_data: 原始表格数据
            expected_columns: 期望的列名列表
        
        Returns:
            调整后的表格数据
        """
        if not table_data:
            return table_data
        
        adjusted_data = []
        for row in table_data:
            adjusted_row = {}
            for col in expected_columns:
                # 如果数据中有该列，使用原值；否则使用空字符串
                adjusted_row[col] = row.get(col, '')
            adjusted_data.append(adjusted_row)
        
        return adjusted_data




