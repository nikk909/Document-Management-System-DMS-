"""
图片处理器
处理图片的插入，支持 Base64 和本地路径
"""
import base64
from pathlib import Path
from typing import Union, Optional, Any
from docx.shared import Inches
from PIL import Image
import io
import sys
import os


class ImageProcessor:
    """
    图片处理器
    负责处理图片的插入，支持 Base64 编码和本地文件路径
    """
    
    def __init__(self):
        """初始化图片处理器"""
        pass
    
    def load_image(self, image_source: Union[str, Path, dict]) -> bytes:
        """
        加载图片数据
        支持 Base64 字符串、本地文件路径、图片ID或字典格式
        
        Args:
            image_source: 图片源，可以是：
                - Base64 字符串（以 'data:image' 开头或纯 Base64）
                - 本地文件路径
                - 图片ID格式：'image_id:31' 或 'id:31'
                - 字典格式：{'id': 31} 或 {'id': 31, 'alt': '...'}
                - HTTP/HTTPS URL
        
        Returns:
            图片的字节数据
        
        Raises:
            FileNotFoundError: 如果文件不存在
            ValueError: 如果格式不支持
        """
        # 如果是字典格式，提取ID
        if isinstance(image_source, dict):
            image_id = image_source.get('id')
            if image_id is not None:
                return self._load_image_by_id(image_id)
            # 如果字典中没有id，尝试使用src字段
            image_source = image_source.get('src', image_source)
        
        # 如果是 Base64 字符串
        if isinstance(image_source, str):
            # 检查是否是图片ID格式：image_id:31 或 id:31
            if image_source.startswith('image_id:') or image_source.startswith('id:'):
                try:
                    # 提取ID
                    if image_source.startswith('image_id:'):
                        image_id = int(image_source[9:])  # 去掉 'image_id:' 前缀
                    else:
                        image_id = int(image_source[3:])  # 去掉 'id:' 前缀
                    return self._load_image_by_id(image_id)
                except ValueError:
                    raise ValueError(f"无效的图片ID格式: {image_source}")
            
            # 检查是否是 data URI 格式
            if image_source.startswith('data:image'):
                # 提取 Base64 部分
                base64_str = image_source.split(',')[1]
                return base64.b64decode(base64_str)
            elif image_source.startswith('base64:'):
                # 支持 base64: 前缀格式
                base64_str = image_source[7:]  # 去掉 'base64:' 前缀
                return base64.b64decode(base64_str)
            elif image_source.startswith('base64,'):
                # 支持 base64, 前缀格式（逗号分隔）
                base64_str = image_source[7:]  # 去掉 'base64,' 前缀
                return base64.b64decode(base64_str)
            elif self._is_base64(image_source):
                # 纯 Base64 字符串
                return base64.b64decode(image_source)
            elif image_source.startswith('http://') or image_source.startswith('https://'):
                # HTTP/HTTPS URL，尝试下载
                # 如果是API URL，只能通过存储管理器获取（避免认证问题）
                if '/api/images/' in image_source:
                    print(f"[DEBUG ImageProcessor] 检测到API URL，通过存储管理器获取: {image_source}", file=sys.stderr, flush=True)
                    # 解析图片ID（在try块外，以便在错误信息中使用）
                    import re
                    match = re.search(r'/api/images/(\d+)/download', image_source)
                    image_id = None
                    if match:
                        image_id = int(match.group(1))
                    
                    try:
                        if not match:
                            raise ValueError(f"无法解析图片URL: {image_source}")
                        
                        # 从数据库获取图片信息
                        # 修复导入路径：从当前文件位置推断backend目录
                        from pathlib import Path as PathLib
                        current_file = PathLib(__file__).resolve()
                        # backend/src/processors/image_processor.py -> backend
                        backend_dir = current_file.parent.parent.parent
                        if str(backend_dir) not in sys.path:
                            sys.path.insert(0, str(backend_dir))
                        
                        from src.storage.storage_manager import StorageManager
                        from src.storage.database import get_db_session, DocumentMetadata
                        
                        # 配置文件路径：项目根目录/config/config.yaml
                        project_root = backend_dir.parent
                        config_path = project_root / "config" / "config.yaml"
                        if not config_path.exists():
                            # 尝试backend/config/config.yaml
                            config_path = backend_dir / "config" / "config.yaml"
                        with get_db_session(config_path=str(config_path)) as db:
                            doc = db.query(DocumentMetadata).filter(
                                DocumentMetadata.id == image_id,
                                DocumentMetadata.category == 'images'
                            ).first()
                            if doc:
                                storage = StorageManager(config_path=str(config_path))
                                image_data = storage.download_bytes(
                                    path=doc.minio_path,
                                    bucket=doc.bucket,
                                    version_id=doc.version_id,
                                    user='system',
                                    user_role='admin',
                                    user_department='system'
                                )
                                print(f"[DEBUG ImageProcessor] 通过存储管理器获取图片成功，大小: {len(image_data)} 字节", file=sys.stderr, flush=True)
                                return image_data
                            else:
                                raise ValueError(f"图片ID {image_id} 在数据库中不存在（category='images'）")
                    except Exception as e:
                        # API URL只能通过存储管理器获取，不能fallback到HTTP下载（因为也需要认证）
                        error_msg = f"无法通过存储管理器获取API图片: {image_source}"
                        error_msg += f"\n错误详情: {str(e)}"
                        error_msg += f"\n提示: API URL只能通过存储管理器获取，请检查："
                        error_msg += f"\n  1. 数据库配置是否正确"
                        if image_id is not None:
                            error_msg += f"\n  2. 图片ID {image_id} 是否存在于数据库中（category='images'）"
                        else:
                            error_msg += f"\n  2. URL格式是否正确（应为 /api/images/<id>/download）"
                        error_msg += f"\n  3. MinIO存储配置是否正确"
                        raise ValueError(error_msg)
                
                # 普通HTTP/HTTPS URL（非API URL），尝试直接下载
                try:
                    import requests
                    print(f"[DEBUG ImageProcessor] 尝试从URL下载图片: {image_source}", file=sys.stderr, flush=True)
                    response = requests.get(image_source, timeout=30, allow_redirects=True)
                    response.raise_for_status()
                    # 检查响应内容类型
                    content_type = response.headers.get('Content-Type', '')
                    if 'image' in content_type or len(response.content) > 0:
                        print(f"[DEBUG ImageProcessor] 成功从URL下载图片，大小: {len(response.content)} 字节", file=sys.stderr, flush=True)
                        return response.content
                    else:
                        raise ValueError(f"URL返回的不是图片内容: {content_type}")
                except Exception as e:
                    print(f"[WARNING] 无法从URL下载图片: {e}", file=sys.stderr, flush=True)
                    raise ValueError(f"无法从URL下载图片: {image_source}\n错误详情: {str(e)}")
            else:
                # 作为文件路径处理
                # 但首先检查：如果image_source仍然是URL格式，说明前面的处理有问题
                if isinstance(image_source, str) and (image_source.startswith('http://') or image_source.startswith('https://')):
                    # 这不应该发生，说明URL处理逻辑有问题
                    raise ValueError(f"URL处理失败，但代码继续执行到文件路径处理: {image_source}")
                
                # 处理以 / 开头的路径（去掉开头的 /）
                if isinstance(image_source, str) and image_source.startswith('/'):
                    image_source = image_source[1:]
                image_source = Path(image_source)
        
        # 如果是路径，读取文件
        # 注意：image_source 可能是 Path 对象或字符串
        if isinstance(image_source, Path):
            image_path = image_source
        else:
            image_path = Path(image_source)
        
        # 如果路径不存在，尝试多个可能的路径
        if not image_path.exists() or not image_path.is_file():
            # 尝试路径列表
            possible_paths = []
            
            # 1. 相对于当前工作目录
            possible_paths.append(Path.cwd() / image_path)
            
            # 2. 如果当前在 backend 目录，尝试项目根目录
            cwd = Path.cwd()
            if cwd.name == 'backend' or 'backend' in str(cwd):
                # 在 backend 目录，向上找到项目根目录
                project_root = cwd.parent
                possible_paths.append(project_root / image_path)
            
            # 3. 如果路径以 image/ 开头，尝试从项目根目录查找
            image_path_str = str(image_path).replace('\\', '/')
            if image_path_str.startswith('image') or image_path_str.startswith('image/'):
                # 尝试从当前目录
                possible_paths.append(Path.cwd() / image_path)
                # 尝试从项目根目录（如果当前在 backend）
                if cwd.name == 'backend' or 'backend' in str(cwd):
                    possible_paths.append(cwd.parent / image_path)
                # 尝试从父目录
                possible_paths.append(cwd.parent / image_path)
                # 尝试从 __file__ 推断的项目根目录
                try:
                    current_file = Path(__file__).resolve()
                    inferred_root = current_file.parent.parent.parent.parent
                    if inferred_root.exists():
                        possible_paths.append(inferred_root / image_path)
                except:
                    pass
            
            # 4. 尝试绝对路径（如果 image_source 是相对路径）
            if not image_path.is_absolute():
                # 尝试从多个可能的根目录查找
                for root in [Path.cwd(), Path.cwd().parent]:
                    possible_paths.append(root / image_path)
            
            # 尝试所有可能的路径
            found = False
            for alt_path in possible_paths:
                try:
                    if alt_path.exists() and alt_path.is_file():
                        image_path = alt_path
                        found = True
                        print(f"[DEBUG ImageProcessor] 找到图片: {image_path}", file=sys.stderr, flush=True)
                        break
                except Exception as e:
                    continue
            
            if not found:
                # 所有路径都尝试过了，打印调试信息
                print(f"[DEBUG ImageProcessor] 当前工作目录: {Path.cwd()}")
                print(f"[DEBUG ImageProcessor] 图片源: {image_source}")
                print(f"[DEBUG ImageProcessor] 已尝试的路径:")
                for p in possible_paths[:10]:  # 显示前10个
                    exists = p.exists()
                    print(f"  - {p} (存在: {exists})")
                
                # 尝试查找项目根目录（包含 image 目录的目录）
                # 方法1: 从当前工作目录向上查找
                current = Path.cwd()
                for _ in range(5):  # 最多向上5级
                    image_dir = current / "image"
                    if image_dir.exists() and image_dir.is_dir():
                        project_root = current
                        # 如果 image_path 是相对路径且以 image/ 开头，直接拼接
                        if str(image_path).startswith('image'):
                            final_path = project_root / image_path
                        else:
                            final_path = project_root / "image" / image_path.name
                        if final_path.exists() and final_path.is_file():
                            print(f"[DEBUG ImageProcessor] 找到图片（方法1）: {final_path}")
                            image_path = final_path
                            found = True
                            break
                    if current == current.parent:  # 到达根目录
                        break
                    current = current.parent
                
                # 方法2: 从 __file__ 位置查找（最可靠的方法）
                if not found:
                    try:
                        # 获取当前文件所在目录
                        current_file = Path(__file__).resolve()
                        # 当前文件应该在 backend/src/processors/image_processor.py
                        # 向上3级到项目根目录（backend -> final_work2）
                        possible_root = current_file.parent.parent.parent.parent
                        image_dir = possible_root / "image"
                        if image_dir.exists() and image_dir.is_dir():
                            # 处理 image/test3.jpg 格式
                            image_path_str = str(image_path).replace('\\', '/')  # 统一使用正斜杠
                            if image_path_str.startswith('image'):
                                final_path = possible_root / image_path_str
                            else:
                                final_path = possible_root / "image" / image_path.name
                            
                            print(f"[DEBUG ImageProcessor] 方法2尝试: {final_path} (项目根: {possible_root}, image_dir存在: {image_dir.exists()})", file=sys.stderr, flush=True)
                            if final_path.exists() and final_path.is_file():
                                print(f"[DEBUG ImageProcessor] 找到图片（方法2）: {final_path}", file=sys.stderr, flush=True)
                                image_path = final_path
                                found = True
                            else:
                                print(f"[DEBUG ImageProcessor] 方法2: 路径不存在 {final_path} (存在: {final_path.exists()}, 是文件: {final_path.is_file() if final_path.exists() else False})", file=sys.stderr, flush=True)
                    except Exception as e:
                        print(f"[DEBUG ImageProcessor] 方法2失败: {e}", file=sys.stderr, flush=True)
                
                # 方法3: 从已知的项目路径查找
                if not found:
                    # 尝试已知的项目根目录路径
                    known_project_roots = [
                        Path("D:/code/AI_trainging/final_work2"),  # 已知的项目路径
                    ]
                    # 如果当前在 backend 目录，向上找到项目根目录
                    cwd_str = str(Path.cwd())
                    if "backend" in cwd_str:
                        backend_idx = cwd_str.find("backend")
                        possible_root = Path(cwd_str[:backend_idx].rstrip("\\/"))
                        if possible_root.exists():
                            known_project_roots.insert(0, possible_root)
                    
                    # 也尝试从 __file__ 推断项目根目录
                    try:
                        current_file = Path(__file__).resolve()
                        # backend/src/processors/image_processor.py -> 向上3级到项目根目录
                        inferred_root = current_file.parent.parent.parent.parent
                        if inferred_root.exists() and (inferred_root / "image").exists():
                            known_project_roots.insert(0, inferred_root)
                    except:
                        pass
                    
                    for root in known_project_roots:
                        if root.exists():
                            # 处理 image/test3.jpg 格式
                            image_path_str = str(image_path)
                            if image_path_str.startswith('image'):
                                final_path = root / image_path_str
                            elif image_path_str.startswith('image/'):
                                final_path = root / image_path_str
                            else:
                                final_path = root / "image" / image_path.name
                            
                            if final_path.exists() and final_path.is_file():
                                print(f"[DEBUG ImageProcessor] 找到图片（方法3）: {final_path}", file=sys.stderr, flush=True)
                                image_path = final_path
                                found = True
                                break
                        else:
                                print(f"[DEBUG ImageProcessor] 方法3尝试: {final_path} (存在: {final_path.exists()}, 是文件: {final_path.is_file() if final_path.exists() else False})", file=sys.stderr, flush=True)
                
                if not found:
                    # 所有路径都尝试过了，抛出错误
                    tried_paths = [str(p) for p in possible_paths[:5]]  # 只显示前5个
                    raise FileNotFoundError(
                        f"图片文件不存在: {image_source}\n"
                        f"当前工作目录: {Path.cwd()}\n"
                        f"已尝试的路径:\n" + "\n".join(f"  - {p}" for p in tried_paths)
                    )
        
        return image_path.read_bytes()
    
    def _load_image_by_id(self, image_id: int) -> bytes:
        """
        通过图片ID从数据库和MinIO加载图片
        
        Args:
            image_id: 图片ID
        
        Returns:
            图片的字节数据
        
        Raises:
            ValueError: 如果图片不存在或加载失败
        """
        try:
            # 从数据库获取图片信息
            from pathlib import Path as PathLib
            current_file = PathLib(__file__).resolve()
            backend_dir = current_file.parent.parent.parent
            if str(backend_dir) not in sys.path:
                sys.path.insert(0, str(backend_dir))
            
            from src.storage.storage_manager import StorageManager
            from src.storage.database import get_db_session, DocumentMetadata
            
            # 配置文件路径
            project_root = backend_dir.parent
            config_path = project_root / "config" / "config.yaml"
            if not config_path.exists():
                config_path = backend_dir / "config" / "config.yaml"
            
            with get_db_session(config_path=str(config_path)) as db:
                doc = db.query(DocumentMetadata).filter(
                    DocumentMetadata.id == image_id,
                    DocumentMetadata.category == 'images'
                ).first()
                if doc:
                    storage = StorageManager(config_path=str(config_path))
                    image_data = storage.download_bytes(
                        path=doc.minio_path,
                        bucket=doc.bucket,
                        version_id=doc.version_id,
                        user='system',
                        user_role='admin',
                        user_department='system'
                    )
                    print(f"[DEBUG ImageProcessor] 通过ID {image_id} 获取图片成功，大小: {len(image_data)} 字节", file=sys.stderr, flush=True)
                    return image_data
                else:
                    raise ValueError(f"图片ID {image_id} 在数据库中不存在（category='images'）")
        except Exception as e:
            error_msg = f"无法通过ID加载图片: {image_id}"
            error_msg += f"\n错误详情: {str(e)}"
            raise ValueError(error_msg)
    
    def get_image_size(self, image_data: bytes) -> tuple:
        """
        获取图片尺寸
        
        Args:
            image_data: 图片字节数据
        
        Returns:
            (width, height) 元组
        """
        try:
            image = Image.open(io.BytesIO(image_data))
            return image.size
        except Exception as e:
            print(f"获取图片尺寸失败: {e}")
            return (0, 0)
    
    def process_for_word(
        self,
        doc: Any,  # Document from python-docx
        placeholder: str,
        image_source: Union[str, Path],
        width: Optional[float] = None,
        height: Optional[float] = None
    ) -> bool:
        """
        为 Word 文档处理图片
        查找占位符 {{image:name}}，替换为图片
        
        Args:
            doc: Word 文档对象
            placeholder: 占位符名称
            image_source: 图片源（Base64 或路径）
            width: 图片宽度（英寸，可选）
            height: 图片高度（英寸，可选）
        
        Returns:
            是否成功处理
        """
        try:
            # 加载图片数据
            image_data = self.load_image(image_source)
            
            # 如果指定了尺寸，使用指定尺寸；否则保持原始比例
            if width is None and height is None:
                # 获取原始尺寸，转换为英寸（假设 96 DPI）
                img_size = self.get_image_size(image_data)
                width = Inches(img_size[0] / 96) if img_size[0] > 0 else Inches(4)
            
            # 查找占位符段落
            placeholder_text = f"{{{{image:{placeholder}}}}}"
            for paragraph in doc.paragraphs:
                if placeholder_text in paragraph.text:
                    # 清除占位符文本并添加图片
                    paragraph.clear()
                    run = paragraph.add_run()
                    # 将图片数据写入临时文件
                    temp_path = Path.cwd() / f"temp_image_{placeholder}.png"
                    temp_path.write_bytes(image_data)
                    try:
                        if width and height:
                            run.add_picture(str(temp_path), width=Inches(width), height=Inches(height))
                        elif width:
                            run.add_picture(str(temp_path), width=Inches(width))
                        elif height:
                            run.add_picture(str(temp_path), height=Inches(height))
                        else:
                            run.add_picture(str(temp_path), width=Inches(4))
                    finally:
                        if temp_path.exists():
                            temp_path.unlink()
                    return True
            
            # 如果在正文中没找到，检查表格单元格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if placeholder_text in paragraph.text:
                                paragraph.clear()
                                run = paragraph.add_run()
                                temp_path = Path.cwd() / f"temp_image_{placeholder}.png"
                                temp_path.write_bytes(image_data)
                                try:
                                    # 表格中的图片尺寸缩小一点
                                    run.add_picture(str(temp_path), width=Inches(2))
                                finally:
                                    if temp_path.exists():
                                        temp_path.unlink()
                                return True
            
            return False
        except Exception as e:
            print(f"处理图片时出错: {e}")
            return False
    
    def process_for_html(
        self,
        template_content: str,
        placeholder: str,
        image_source: Union[str, Path]
    ) -> str:
        """
        为 HTML/PDF 模板处理图片
        将占位符替换为图片标签
        
        Args:
            template_content: HTML 模板内容
            placeholder: 占位符名称
            image_source: 图片源（Base64 或路径）
        
        Returns:
            处理后的 HTML 内容
        """
        try:
            # 加载图片数据
            image_data = self.load_image(image_source)
            
            # 转换为 Base64
            base64_str = base64.b64encode(image_data).decode('utf-8')
            
            # 检测图片格式（简单检测，默认 PNG）
            img_format = 'png'
            if isinstance(image_source, Path):
                ext = image_source.suffix.lower()
                if ext in ['.jpg', '.jpeg']:
                    img_format = 'jpeg'
                elif ext == '.gif':
                    img_format = 'gif'
            
            # 生成图片标签
            img_tag = f'<img src="data:image/{img_format};base64,{base64_str}" alt="{placeholder}" style="max-width: 100%; height: auto;" />'
            
            # 替换占位符
            return template_content.replace(f"{{{{image:{placeholder}}}}}", img_tag)
        except Exception as e:
            print(f"处理图片时出错: {e}")
            return template_content.replace(f"{{{{image:{placeholder}}}}}", "")
    
    def _is_base64(self, s: str) -> bool:
        """
        检查字符串是否是 Base64 编码
        
        Args:
            s: 待检查的字符串
        
        Returns:
            是否是 Base64 编码
        """
        try:
            # 尝试解码
            base64.b64decode(s, validate=True)
            return True
        except Exception:
            return False

