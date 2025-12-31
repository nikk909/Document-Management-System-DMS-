"""
图表处理器
基于数据生成折线图/柱状图，并嵌入文档
"""
import matplotlib
matplotlib.use('Agg')  # 使用非交互式后端
import matplotlib.pyplot as plt
from pathlib import Path
from typing import Dict, Any, List, Optional
import base64
import io
from datetime import datetime


class ChartProcessor:
    """
    图表处理器
    负责生成图表（折线图/柱状图），并转换为图片嵌入文档
    """
    
    def __init__(self, temp_dir: Optional[Path] = None):
        """
        初始化图表处理器
        
        Args:
            temp_dir: 临时文件目录（用于保存图表图片）
        """
        self.temp_dir = temp_dir or Path.cwd() / "temp_charts"
        self.temp_dir.mkdir(exist_ok=True)
    
    def generate_chart(
        self,
        chart_data: Dict[str, Any],
        chart_type: str = 'line',
        output_path: Optional[Path] = None
    ) -> Path:
        """
        生成图表并保存为图片
        
        Args:
            chart_data: 图表数据字典，格式：
                {
                    'data': [
                        {'x': 'Jan', 'y': 100},
                        {'x': 'Feb', 'y': 200},
                        ...
                    ]
                    或
                    'x': ['Jan', 'Feb', ...],
                    'y': [100, 200, ...]
                }
            chart_type: 图表类型（'line' 折线图 / 'bar' 柱状图）
            output_path: 输出路径（可选，默认自动生成）
        
        Returns:
            生成的图片文件路径
        """
        # 解析数据（支持多种格式）
        x_values = []
        y_values = []
        
        # 格式1: series和labels格式（新格式）
        if 'series' in chart_data and 'labels' in chart_data:
            labels = chart_data['labels']
            series_list = chart_data['series']
            if series_list and len(series_list) > 0:
                first_series = series_list[0]
                if 'points' in first_series:
                    y_values = first_series['points']
                    x_values = labels if len(labels) == len(y_values) else [str(i+1) for i in range(len(y_values))]
                elif 'data' in first_series:
                    y_values = first_series['data']
                    x_values = labels if len(labels) == len(y_values) else [str(i+1) for i in range(len(y_values))]
        
        # 格式2: data格式 {'data': [{'x': ..., 'y': ...}, ...]}
        if not x_values and 'data' in chart_data:
            data_list = chart_data['data']
            x_values = [item.get('x', item.get('label', '')) for item in data_list]
            y_values = [item.get('y', item.get('value', 0)) for item in data_list]
        
        # 格式3: x和y格式 {'x': [...], 'y': [...]}
        if not x_values:
            x_values = chart_data.get('x', [])
            y_values = chart_data.get('y', [])
        
        if not x_values or not y_values:
            raise ValueError("图表数据为空")
        
        # 创建图表
        plt.figure(figsize=(10, 6))
        
        if chart_type == 'line':
            plt.plot(x_values, y_values, marker='o', linewidth=2, markersize=8)
        elif chart_type == 'bar':
            plt.bar(x_values, y_values, width=0.6)
        else:
            raise ValueError(f"不支持的图表类型: {chart_type}")
        
        # 设置标题和标签
        title = chart_data.get('title', f'{chart_type.title()} Chart')
        x_label = chart_data.get('x_label', 'X Axis')
        y_label = chart_data.get('y_label', 'Y Axis')
        
        plt.title(title, fontsize=14, fontweight='bold')
        plt.xlabel(x_label, fontsize=12)
        plt.ylabel(y_label, fontsize=12)
        
        # 添加网格
        plt.grid(True, alpha=0.3)
        
        # 旋转 x 轴标签（如果太长）
        plt.xticks(rotation=45, ha='right')
        
        # 调整布局
        plt.tight_layout()
        
        # 保存图片
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = self.temp_dir / f"chart_{timestamp}.png"
        
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()  # 关闭图形，释放内存
        
        return output_path
    
    def generate_chart_base64(
        self,
        chart_data: Dict[str, Any],
        chart_type: str = 'line'
    ) -> str:
        """
        生成图表并转换为 Base64 编码字符串
        
        Args:
            chart_data: 图表数据
            chart_type: 图表类型
        
        Returns:
            Base64 编码的图片字符串
        """
        # 生成临时图片
        temp_path = self.generate_chart(chart_data, chart_type)
        
        # 读取图片并转换为 Base64
        with open(temp_path, 'rb') as f:
            image_data = f.read()
            base64_str = base64.b64encode(image_data).decode('utf-8')
        
        # 删除临时文件
        temp_path.unlink()
        
        return base64_str
    
    def process_for_word(
        self,
        doc: Any,  # Document from python-docx
        placeholder: str,
        chart_data: Dict[str, Any],
        chart_type: str = 'line'
    ) -> bool:
        """
        为 Word 文档处理图表
        查找占位符 {{chart:name}}，替换为图表图片
        
        Args:
            doc: Word 文档对象
            placeholder: 占位符名称
            chart_data: 图表数据
            chart_type: 图表类型
        
        Returns:
            是否成功处理
        """
        try:
            from docx.shared import Inches
            
            # 生成图表图片
            chart_path = self.generate_chart(chart_data, chart_type)
            
            # 查找占位符段落
            placeholder_text = f"{{{{chart:{placeholder}}}}}"
            for paragraph in doc.paragraphs:
                if placeholder_text in paragraph.text:
                    # 清除占位符文本并添加图片
                    paragraph.clear()
                    run = paragraph.add_run()
                    run.add_picture(str(chart_path), width=Inches(6))
                    chart_path.unlink()
                    return True
            
            # 如果在正文中没找到，检查表格单元格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if placeholder_text in paragraph.text:
                                paragraph.clear()
                                run = paragraph.add_run()
                                # 表格中的图表尺寸缩小一点
                                run.add_picture(str(chart_path), width=Inches(4))
                                chart_path.unlink()
                                return True
            
            # 如果没找到占位符，删除临时文件
            if chart_path.exists():
                chart_path.unlink()
            
            return False
        except Exception as e:
            print(f"处理图表时出错: {e}")
            return False
    
    def process_for_html(
        self,
        template_content: str,
        placeholder: str,
        chart_data: Dict[str, Any],
        chart_type: str = 'line'
    ) -> str:
        """
        为 HTML/PDF 模板处理图表
        将占位符替换为图片标签
        
        Args:
            template_content: HTML 模板内容
            placeholder: 占位符名称
            chart_data: 图表数据
            chart_type: 图表类型
        
        Returns:
            处理后的 HTML 内容
        """
        try:
            # 生成图表并转换为 Base64
            base64_str = self.generate_chart_base64(chart_data, chart_type)
            
            # 生成图片标签
            img_tag = f'<img src="data:image/png;base64,{base64_str}" alt="Chart" style="max-width: 100%; height: auto;" />'
            
            # 替换占位符
            return template_content.replace(f"{{{{chart:{placeholder}}}}}", img_tag)
        except Exception as e:
            print(f"处理图表时出错: {e}")
            return template_content.replace(f"{{{{chart:{placeholder}}}}}", "")




