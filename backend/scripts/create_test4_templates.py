# -*- coding: utf-8 -*-
"""创建test4的Word模板"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create CSV Word template
doc = Document()

# Title
title = doc.add_heading('员工信息表', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Info
doc.add_paragraph('生成时间: {{ now.strftime("%Y-%m-%d %H:%M:%S") }}')

# Table placeholder
doc.add_paragraph('{% if table_data %}')
doc.add_paragraph('{{table:data}}')
doc.add_paragraph('{% endif %}')

# Footer
footer = doc.add_paragraph('本文档由系统自动生成')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save('testdata/template/test4/csv/test4.docx')
print('CSV Word template created: testdata/template/test4/csv/test4.docx')

# Create JSON Word template
doc2 = Document()

# Title with jinja2
title2 = doc2.add_heading('{{ document.title if document else "员工信息表" }}', 0)
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Meta info
doc2.add_paragraph('{% if document %}')
doc2.add_paragraph('日期: {{ document.date }} | 部门: {{ document.department }}')
doc2.add_paragraph('{% endif %}')

# Content
doc2.add_paragraph('{% if content %}')
doc2.add_paragraph('{{ content }}')
doc2.add_paragraph('{% endif %}')

# Table
doc2.add_heading('员工数据', level=1)
doc2.add_paragraph('{% if table_data %}')
doc2.add_paragraph('{{table:data}}')
doc2.add_paragraph('{% endif %}')

# Footer
footer2 = doc2.add_paragraph('生成时间: {{ now.strftime("%Y-%m-%d %H:%M:%S") }}')
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc2.save('testdata/template/test4/json/test4.docx')
print('JSON Word template created: testdata/template/test4/json/test4.docx')
